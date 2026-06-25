from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_dat_docx import (
    BLUE,
    DARK_BLUE,
    MUTED,
    add_bullet_list,
    add_header_footer,
    add_heading,
    add_paragraphs,
    add_table,
    configure_section,
    create_styles,
    ensure_font,
    set_cell_margins,
    set_cell_shading,
    set_column_widths,
    set_table_borders,
    set_table_width,
)


ROOT = Path(r"C:\Users\Arthur\Pentest")
CODE_PROMO = "MSI-4-27-CS-C-ISI-PARIS"
PROJECT_NAME = "Reconforge"


INDIVIDUALS = [
    {
        "first_name": "Arthur",
        "last_name": "FLAMENT",
        "role": "Interface, reporting, documentation et recette",
        "hours": "46 j.h",
        "focus_table": [
            ("Interface web", "Conception de l'experience operateur, organisation des modules et logique de saisie cible."),
            ("Reporting", "Consolidation des sorties en JSON/PDF et lisibilite des artefacts pour la soutenance."),
            ("Documentation", "README, DAT, coherence documentaire et formalisation des choix techniques."),
            ("Recette", "Verification fonctionnelle des parcours utilisateurs et preparation des demonstrations."),
        ],
        "future_text": (
            "Du point de vue de mon perimetre, l'avenir de Reconforge passe d'abord par une amelioration de l'experience "
            "operateur. L'interface actuelle est volontairement sobre et efficace, mais elle peut encore gagner en accompagnement "
            "contextuel. Je vois comme priorites l'ajout d'aides par module, de messages de validation plus explicites, d'une "
            "recherche plus riche dans l'historique et d'une meilleure visualisation des sessions multi-outils. Dans le meme "
            "mouvement, la partie reporting merite une montee en gamme importante : gabarits plus soignes, separation plus claire "
            "entre synthese executive et detail technique, et possibilite future d'exports HTML ou DOCX en plus du PDF."
        ),
        "future_text_2": (
            "Je pense aussi que l'avenir de la solution repose sur une meilleure articulation entre interface, persistance et "
            "restitution. Aujourd'hui, l'historique reconstruit une vue a partir des artefacts locaux. Demain, un branchement plus "
            "complet vers PostgreSQL et MinIO permettrait de proposer de vrais filtres multicriteres, un suivi par utilisateur, des "
            "etats de validation de rapports et une logique d'archivage plus mature. Cela renforcerait la valeur percue du produit : "
            "on passerait d'une toolbox de demonstration reussie a une base plus proche d'une vraie plateforme d'orchestration."
        ),
        "limits_text": (
            "La principale limite technique que j'ai rencontree sur mon perimetre concerne le compromis entre simplicite de l'interface "
            "et richesse fonctionnelle des outils integres. Exposer trop d'options aurait rendu la page illisible et peu defendable en "
            "soutenance ; en exposer trop peu pouvait donner l'impression d'une abstraction excessive. Il a donc fallu trouver un point "
            "d'equilibre, ce qui a parfois implique de renoncer a certaines options natives des outils pour privilegier la coherence globale."
        ),
        "limits_text_2": (
            "La seconde limite touche au reporting. Produire des artefacts JSON et PDF coherents et presentables est deja une vraie valeur, "
            "mais la restitution reste encore sobre. Cela tient au fait que le projet a d'abord privilegie la fiabilite du flux plutot que la "
            "richesse visuelle du document final. J'identifie aussi une limite sur la recette visuelle complete du projet : lorsque l'interface, "
            "les rapports et les donnees evoluent en meme temps, il devient necessaire d'introduire davantage de tests de non-regression et de "
            "rituels de verification systematiques."
        ),
        "user_doc": [
            "Lancement : copier .env.example vers .env, verifier Docker Desktop, puis lancer docker compose up --build -d depuis la racine du depot.",
            "Acces : ouvrir http://localhost:8000, puis se connecter si l'authentification locale est activee.",
            "Utilisation : saisir une cible, cocher un ou plusieurs modules, lancer l'execution et suivre la console temps reel.",
            "Resultats : consulter l'historique a droite et telecharger les artefacts JSON/PDF de chaque session.",
            "Demonstration : pour un lab web, demarrer aussi Juice Shop via le profil lab et utiliser une URL complete comme cible.",
        ],
        "challenges": (
            "Le defi principal pour moi a ete de rendre la plateforme intelligible sans la rendre artificiellement simple. Il ne suffisait pas "
            "de faire une belle interface ; il fallait que cette interface reste fidele au comportement reel des outils, tout en etant comprehensible "
            "par un evaluateur ou un utilisateur qui ne vit pas dans la ligne de commande. J'ai aussi du maintenir une coherence forte entre l'ergonomie, "
            "le storytelling du projet, le DAT, le README et la demo orale."
        ),
        "strengths": [
            "Capacite a structurer une interface lisible a partir d'exigences techniques heterogenes.",
            "Sens de la restitution : reporting, documentation et clarte de presentation.",
            "Vision transversale du projet, utile pour garder une coherence entre technique et livrables.",
        ],
        "weaknesses": [
            "Tendance a vouloir trop polir la forme avant que tous les mecanismes techniques soient completement stabilises.",
            "Besoin de renforcer encore l'automatisation de mes verifications visuelles et fonctionnelles.",
            "Marge de progression sur la formalisation tres en amont des criteres UX et de recette.",
        ],
        "skills": [
            "Structuration d'une interface FastAPI/Jinja2 sobre et demonstrative.",
            "Conception de parcours utilisateur adaptes a des outils de pentest.",
            "Generation et mise en coherence d'artefacts de reporting.",
            "Documentation technique et projet a un niveau defendable en soutenance.",
        ],
        "improvements": [
            "Mettre en place des tests de non-regression plus systematiques sur l'interface et le reporting.",
            "Approfondir les pratiques de design UX appliquees aux outils techniques.",
            "Mieux anticiper les besoins de visualisation des donnees des la conception du backend.",
        ],
    },
    {
        "first_name": "Mikail",
        "last_name": "GULRAIZ",
        "role": "Backend, execution, securite et orchestration",
        "hours": "50 j.h",
        "focus_table": [
            ("Backend coeur", "Structuration des services, logique d'execution commune et cohesion technique de l'application."),
            ("Execution des modules", "Integration des outils, normalisation des cibles, timeouts et streaming des sorties."),
            ("SSH Kali", "Delegation d'execution vers la VM Kali et gestion de la separation entre cockpit web et environnement operateur."),
            ("Securite", "Authentification locale, A2F TOTP, audit log et cadrage des limites de durcissement."),
        ],
        "future_text": (
            "Depuis mon perimetre backend et securite, l'avenir de Reconforge passe d'abord par une industrialisation plus forte de l'execution. "
            "Aujourd'hui, le mode local et le mode SSH couvrent deja le besoin de demonstration, mais une evolution naturelle consisterait a mieux "
            "separer les profils d'execution : mode local simplifie, mode Kali SSH complet, puis mode asynchrone pour les scans longs. L'activation "
            "plus poussee de Celery et Redis permettrait alors de sortir d'un flux purement synchrone et d'aller vers des campagnes plus longues, "
            "plus resilientes et plus faciles a reprendre."
        ),
        "future_text_2": (
            "Je vois aussi comme priorite majeure la maturation de la couche de securite et de persistance. PostgreSQL pourrait devenir la vraie source "
            "de verite pour les comptes, les roles, les sessions, les scans et les audit events. MinIO prendrait alors en charge les artefacts, tandis "
            "qu'un RBAC minimal, un chiffrement plus robuste des secrets et une gestion MFA par utilisateur feraient franchir un cap important a la plateforme. "
            "A plus long terme, la formalisation d'un systeme de plugins permettrait d'integrer de nouveaux outils sans alourdir le coeur applicatif."
        ),
        "limits_text": (
            "La limite technique la plus structurante que j'ai rencontree concerne l'heterogeneite des outils integres. Tous n'ont pas la meme syntaxe, "
            "les memes besoins reseau, les memes temps d'execution ni les memes comportements d'erreur. Construire un moteur commun capable d'absorber cette "
            "diversite sans produire un backend fragile a demande un travail important de normalisation, de garde-fous et de messages d'erreur."
        ),
        "limits_text_2": (
            "La deuxieme limite concerne l'environnement d'execution lui-meme. Le choix de s'appuyer sur Kali via SSH est pertinent techniquement, mais il "
            "introduit une dependance reseau et un enjeu de fiabilite supplementaire : disponibilite de la VM, chemins de binaires, wordlists, host keys, "
            "time-outs et ecarts entre environnement conteneurise et environnement operateur. Sur le plan securite, l'authentification locale et l'A2F sont "
            "deja de bonnes bases, mais elles restent encore en dessous de ce qu'exigerait un vrai contexte multi-utilisateurs."
        ),
        "user_doc": [
            "Configuration backend : verifier les variables APP_*, EXECUTION_MODE et les timeouts des modules dans .env.",
            "Mode SSH : renseigner KALI_SSH_HOST, KALI_SSH_PORT, KALI_SSH_USERNAME, KALI_SSH_PASSWORD ou la cle d'acces avant le lancement.",
            "Dependances outils : verifier sur Kali la presence des binaires et des wordlists utilises par Hydra, Gobuster ou theHarvester.",
            "API : utiliser /docs pour verifier les endpoints exposes et /health pour controler le bon demarrage du service.",
            "Tracabilite : consulter storage/logs/audit.log et storage/reports/ pour verifier la persistence des sessions et des artefacts.",
        ],
        "challenges": (
            "Le principal defi personnel a ete de construire une architecture assez generique pour accueillir des outils tres differents, sans tomber dans un "
            "simple assemblage de scripts. Il fallait penser a la fois en termes de securite, de robustesse d'execution, d'ergonomie du backend et de coherence "
            "avec la demo finale. J'ai aussi du arbitrer en permanence entre richesse fonctionnelle et soutenabilite technique, notamment sur l'execution SSH, "
            "les validations d'entrees et la gestion des erreurs."
        ),
        "strengths": [
            "Capacite a structurer un backend modulaire autour d'un contrat d'execution commun.",
            "Rigueur dans le cadrage de l'execution, des validations et des garde-fous.",
            "Bonne lecture des compromis entre securite, demonstration et faisabilite technique.",
        ],
        "weaknesses": [
            "Tendance a pousser la robustesse technique avant la simplification visible pour l'utilisateur final.",
            "Besoin de renforcer encore l'automatisation des tests backend et d'integration.",
            "Marge de progression sur l'anticipation des impacts de certaines decisions techniques sur la restitution finale.",
        ],
        "skills": [
            "Architecture backend modulaire avec FastAPI.",
            "Execution de commandes en local et via SSH avec streaming.",
            "Gestion de la securite d'acces : authentification, A2F et audit.",
            "Conception de services capables d'integrer des outils heterogenes dans un cadre commun.",
        ],
        "improvements": [
            "Renforcer les tests d'integration autour des outils et du mode SSH.",
            "Approfondir les mecanismes de persistance relationnelle et de gouvernance des artefacts.",
            "Formaliser plus tot les profils d'execution et les politiques de securite ciblees.",
        ],
    },
]


def add_custom_cover(document: Document, person: dict[str, object]) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    run = p.add_run("RENDU INDIVIDUEL")
    ensure_font(run, name="Calibri", size=14, color=MUTED, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(PROJECT_NAME)
    ensure_font(run, name="Calibri", size=28, color=DARK_BLUE, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{person['first_name']} {person['last_name']}")
    ensure_font(run, size=18, color=BLUE, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(
        "Perspectives d'evolution, analyse critique, documentation utilisateur et reflexion personnelle "
        "autour du projet Reconforge."
    )
    ensure_font(run, size=11)

    table = document.add_table(rows=5, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_width(table)
    set_table_borders(table)
    set_column_widths(table, [4.0, 11.5])
    metadata = [
        ("Projet", PROJECT_NAME),
        ("Code promo", CODE_PROMO),
        ("Etudiant", f"{person['first_name']} {person['last_name']}"),
        ("Perimetre principal", str(person["role"])),
        ("Date", datetime.now().strftime("%d/%m/%Y")),
    ]

    for row_idx, (label, value) in enumerate(metadata):
        left, right = table.rows[row_idx].cells
        for cell in (left, right):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, "F2F4F7")
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        lrun = lp.add_run(label)
        rrun = rp.add_run(value)
        ensure_font(lrun, size=10, bold=True)
        ensure_font(rrun, size=10)

    document.add_paragraph()
    document.add_page_break()


def add_role_summary(document: Document, person: dict[str, object]) -> None:
    add_heading(document, "1. Positionnement individuel dans le projet", 1)
    add_paragraphs(
        document,
        (
            f"Dans le cadre de Reconforge, {person['first_name']} {person['last_name']} a travaille principalement sur le perimetre "
            f"suivant : {person['role']}. Le volume de charge estime sur ce perimetre est d'environ {person['hours']}. "
            "Cette contribution s'inscrit dans une logique de travail transverse : meme si chaque membre avait un coeur de responsabilite "
            "bien identifie, les arbitrages d'architecture, de planning et de demonstration ont ete construits a deux."
        ),
    )
    add_table(
        document,
        ["Axe de contribution", "Contenu du travail realise"],
        person["focus_table"],
        [4.0, 12.0],
    )


def add_personal_sections(document: Document, person: dict[str, object]) -> None:
    add_heading(document, "2. Perspectives d'evolution et reflexion sur l'avenir de la solution", 1)
    add_paragraphs(document, str(person["future_text"]))
    add_paragraphs(document, str(person["future_text_2"]))

    add_heading(document, "3. Analyse critique sur les limites techniques rencontrees", 1)
    add_paragraphs(document, str(person["limits_text"]))
    add_paragraphs(document, str(person["limits_text_2"]))

    add_heading(document, "4. Annexes", 1)
    add_heading(document, "4.1 Documentation utilisateur", 2)
    add_paragraphs(
        document,
        "Cette annexe rappelle, de maniere concise, les etapes essentielles pour installer, lancer et utiliser la plateforme selon mon perimetre de contribution.",
    )
    add_bullet_list(document, list(person["user_doc"]))

    add_heading(document, "4.2 Analyse personnelle", 2)
    add_heading(document, "4.2.1 Reflexion sur les defis rencontres", 3)
    add_paragraphs(document, str(person["challenges"]))

    add_heading(document, "4.2.2 Identification des forces et faiblesses personnelles", 3)
    add_paragraphs(document, "Forces personnelles identifiees pendant le projet :")
    add_bullet_list(document, list(person["strengths"]))
    add_paragraphs(document, "Faiblesses ou points de vigilance identifies pendant le projet :")
    add_bullet_list(document, list(person["weaknesses"]))

    add_heading(document, "4.2.3 Competences developpees", 3)
    add_bullet_list(document, list(person["skills"]))

    add_heading(document, "4.2.4 Axes d'amelioration personnels pour de futurs projets", 3)
    add_bullet_list(document, list(person["improvements"]))


def build_doc(person: dict[str, object]) -> Path:
    document = Document()
    configure_section(document.sections[0])
    create_styles(document)
    add_header_footer(document.sections[0], first_page=True)

    add_custom_cover(document, person)
    add_role_summary(document, person)
    add_personal_sections(document, person)

    output = ROOT / f"PE-2526_{CODE_PROMO}_{person['first_name']}{person['last_name']}.docx"
    document.save(output)
    return output


def main() -> None:
    for person in INDIVIDUALS:
        path = build_doc(person)
        print(path)


if __name__ == "__main__":
    main()
