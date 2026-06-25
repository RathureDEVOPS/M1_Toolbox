from __future__ import annotations

from datetime import datetime
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(r"C:\Users\Arthur\Pentest\DAT.docx")
APP_NAME = "Reconforge"
APP_TAGLINE = "Plateforme locale et conteneurisee d'orchestration d'outils de pentest"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x25, 0x25, 0x25)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "F2F4F7"
BORDER = "D0D7DE"
PLACEHOLDER = "EEF4FB"


def save_document(document: Document) -> Path:
    candidates = [
        OUTPUT_PATH,
        OUTPUT_PATH.with_name("Reconforge_DAT.docx"),
        OUTPUT_PATH.with_name(f"Reconforge_DAT_{datetime.now():%Y%m%d_%H%M%S}.docx"),
    ]
    last_error: PermissionError | None = None
    for path in candidates:
        try:
            document.save(path)
            return path
        except PermissionError as error:
            last_error = error
    if last_error:
        raise last_error
    return OUTPUT_PATH


REQUIREMENT_ROWS = [
    (
        "Architecture générale",
        "Python 3.x, framework web moderne, architecture modulaire et automatisée.",
        "Couvert",
        "Application Python 3.12, FastAPI, Jinja2, découpage en routes, services, modules et intégrations.",
    ),
    (
        "Modules fonctionnels",
        "Intégrer Nmap, OpenVAS/Nessus, Metasploit, scripts personnalisés.",
        "Partiellement couvert",
        "Nmap est intégré, la structure modulaire est prête pour d'autres outils, mais OpenVAS/Nessus et Metasploit ne sont pas encore branchés.",
    ),
    (
        "Outils open source",
        "Prévoir ZAP, SQLmap, Hydra, Aircrack-ng et outils complémentaires.",
        "Partiellement couvert",
        "SQLmap et Hydra sont déjà disponibles, ainsi que Nikto, SSLyze, WhatWeb, theHarvester, Gobuster et Wireshark. ZAP et Aircrack-ng restent à planifier.",
    ),
    (
        "Stockage et base de données",
        "S'appuyer sur PostgreSQL et MinIO.",
        "Partiellement couvert",
        "Les conteneurs PostgreSQL et MinIO sont présents dans l'architecture, mais la persistance des rapports est encore réalisée sur le système de fichiers local.",
    ),
    (
        "Reporting",
        "Prévoir Jinja2, export PDF/HTML/CSV, D3.js ou Matplotlib.",
        "Partiellement couvert",
        "L'interface repose sur Jinja2 et les rapports sont produits en JSON et PDF. Les exports HTML/CSV dédiés et les graphiques restent une évolution.",
    ),
    (
        "Orchestration",
        "Prévoir Celery et Redis.",
        "Partiellement couvert",
        "Redis et un worker Celery sont définis dans Docker Compose, mais le chemin d'exécution courant reste synchrone pour privilégier la simplicité opérationnelle.",
    ),
    (
        "Sécurité",
        "Chiffrement Fernet, HTTPS, RBAC, audit logs.",
        "Partiellement couvert",
        "Audit logs et authentification avec A2F sont opérationnels. Fernet, RBAC et HTTPS ne sont pas encore industrialisés dans le flux de production.",
    ),
    (
        "Conteneurisation",
        "Docker, Docker Compose, CI/CD GitLab.",
        "Partiellement couvert",
        "Docker et Docker Compose sont en place. La CI/CD GitLab est identifiée comme une étape de maturité future.",
    ),
    (
        "Évolutivité",
        "Système de plugins et API documentée via Swagger.",
        "Partiellement couvert",
        "Swagger est disponible via FastAPI. Le système de plugins n'est pas encore formalisé, mais la structure modulaire a été pensée pour l'accueillir.",
    ),
]


DOCKER_SERVICES = [
    ("web", "Application FastAPI/Jinja2", "Expose l'interface locale, les routes web, l'API et la documentation Swagger."),
    ("worker", "Celery worker", "Prépare l'externalisation des traitements longs vers une exécution asynchrone."),
    ("postgres", "Base relationnelle", "Cible de persistance relationnelle pour les métadonnées et l'historique structuré."),
    ("redis", "Broker et backend", "Support de Celery et mécanisme de coordination pour la montée en charge future."),
    ("minio", "Stockage objet", "Cible prévue pour l'archivage des rapports et artefacts volumineux."),
    ("juice-shop", "Lab d'application web vulnérable", "Cible de démonstration utilisée pour valider les modules orientés web."),
]


MODULE_ROWS = [
    ("Sécurité (SOC, EDR, XDR)", "Nmap", "Reconnaissance réseau contrôlée, détection de services et scripts NSE sur liste blanche."),
    ("Sécurité (SOC, EDR, XDR)", "Wireshark", "Capture réseau courte via tshark avec statistiques et synthèse textuelle."),
    ("Développement SaaS", "WhatWeb", "Fingerprint web rapide des technologies, frameworks et serveurs exposés."),
    ("Développement SaaS", "theHarvester", "Reconnaissance OSINT sur domaine : sous-domaines, hôtes et indices publics."),
    ("Développement SaaS", "Gobuster", "Enumération de contenu web maîtrisée avec temporisation et exclusion automatique de réponses génériques."),
    ("Développement SaaS", "SQLmap", "Détection d'injections SQL sur URL HTTP(S) avec options limitées et exécution non interactive."),
    ("Infrastructure", "Hydra", "Tests d'authentification SSH avec saisie manuelle ou wordlists Kali sélectionnées."),
    ("Support client", "Nikto", "Audit web orienté configuration serveur, contenu exposé et points faibles applicatifs."),
    ("Support client", "SSLyze", "Vérification de la surface TLS et des paramètres cryptographiques d'un service."),
    ("RH / Administration", "Auth Audit", "Audit contrôlé d'un formulaire ou endpoint d'authentification sans dériver vers un bruteforce massif."),
]


TEST_SCENARIOS = [
    ("Connexion sécurisée", "Validation du parcours login + A2F", "OK", "Le premier facteur redirige soit vers l'initialisation TOTP, soit vers la vérification du code OTP."),
    ("Exécution Nmap", "Scan simple sur cible IP ou domaine", "OK", "Commande lancée, flux temps réel affiché, rapport JSON/PDF généré."),
    ("Exécution web multi-outils", "WhatWeb + Gobuster + SQLmap + Nikto sur lab", "OK", "Les modules s'enchaînent dans une même session et l'historique n'affiche qu'une entrée agrégée."),
    ("Hydra SSH", "Test ciblé sur service SSH", "OK", "La configuration conditionnelle apparaît uniquement lorsque le module est coché."),
    ("Auth Audit", "Analyse d'un endpoint de connexion", "OK", "Le module restitue les réponses HTTP, les cookies, les protections et un résumé interprétable."),
    ("Mode SSH Kali", "Délégation d'exécution vers la VM Kali", "OK", "Le site reste conteneurisé tandis que les outils s'exécutent dans l'environnement Kali distant."),
    ("Reporting", "Téléchargement JSON/PDF", "OK", "Chaque session dispose d'artefacts cohérents et accessibles via l'interface."),
]


API_ENDPOINTS = [
    ("GET", "/health", "Sonde de vie de l'application."),
    ("GET", "/docs", "Documentation Swagger générée par FastAPI."),
    ("GET", "/api/modules", "Catalogue des modules disponibles dans l'interface."),
    ("POST", "/api/scans/<module>", "Lancement d'un scan unitaire avec réponse complète."),
    ("POST", "/api/scans/<module>/stream", "Lancement d'un scan avec retour progressif en NDJSON."),
    ("POST", "/api/scans/session", "Agrégation d'une session multi-modules et génération des artefacts finaux."),
    ("GET", "/api/scans/history", "Liste des dernières exécutions disponibles dans l'historique."),
    ("GET", "/api/scans/{scan_id}/artifacts/{format}", "Téléchargement d'un rapport JSON ou PDF."),
]


ENV_GROUPS = [
    (
        "Configuration générale",
        [
            "APP_NAME, APP_HOST, APP_PORT",
            "EXECUTION_MODE",
        ],
    ),
    (
        "Authentification et session",
        [
            "TOOLBOX_AUTH_ENABLED",
            "TOOLBOX_AUTH_USERNAME",
            "TOOLBOX_AUTH_PASSWORD",
            "TOOLBOX_AUTH_TOTP_SECRET",
            "TOOLBOX_AUTH_TOTP_ISSUER",
            "AUTH_SESSION_SECRET",
        ],
    ),
    (
        "Connexion Kali",
        [
            "KALI_SSH_HOST",
            "KALI_SSH_PORT",
            "KALI_SSH_USERNAME",
            "KALI_SSH_PASSWORD",
            "KALI_SSH_KEY_PATH",
            "KALI_SSH_ALLOW_UNKNOWN_HOST",
        ],
    ),
    (
        "Modules et timeouts",
        [
            "NMAP_BINARY / NMAP_TIMEOUT",
            "NIKTO_BINARY / NIKTO_TIMEOUT",
            "SQLMAP_BINARY / SQLMAP_TIMEOUT",
            "SSLYZE_BINARY / SSLYZE_TIMEOUT",
            "WHATWEB_BINARY / WHATWEB_TIMEOUT",
            "THEHARVESTER_BINARY / THEHARVESTER_TIMEOUT / THEHARVESTER_SOURCES / THEHARVESTER_LIMIT",
            "GOBUSTER_BINARY / GOBUSTER_TIMEOUT / GOBUSTER_WORDLIST / GOBUSTER_THREADS / GOBUSTER_DELAY",
            "HYDRA_BINARY / HYDRA_TIMEOUT / HYDRA_THREADS / wordlists",
            "AUTH_AUDIT_TIMEOUT / valeurs par défaut",
            "WIRESHARK_BINARY / WIRESHARK_TIMEOUT / WIRESHARK_DEFAULT_INTERFACE",
        ],
    ),
    (
        "Services externes",
        [
            "POSTGRES_DSN",
            "REDIS_URL",
            "MINIO_ENDPOINT",
            "MINIO_ACCESS_KEY",
            "MINIO_SECRET_KEY",
            "MINIO_SECURE",
            "FERNET_KEY",
            "GITHUB_TOKEN",
        ],
    ),
]

PROJECT_TIMELINE_ROWS = [
    ("Cadrage, architecture et backlog", "Gulraiz Mikail + Arthur Flament", "X", "", "", "", "", ""),
    ("Socle Docker, CI/CD et observabilite", "Gulraiz Mikail / Arthur Flament", "X", "X", "", "", "", ""),
    ("Backend coeur et execution commune", "Gulraiz Mikail", "", "X", "X", "", "", ""),
    ("Interface web et experience operateur", "Arthur Flament", "", "X", "X", "", "", ""),
    ("Integration des modules coeur", "Gulraiz Mikail / Arthur Flament", "", "", "X", "X", "", ""),
    ("Kali SSH, lab et stabilisation environnement", "Gulraiz Mikail / Arthur Flament", "", "", "X", "X", "", ""),
    ("Reporting, historique et restitution PDF", "Arthur Flament", "", "", "", "X", "X", ""),
    ("Authentification, A2F et audit", "Gulraiz Mikail / Arthur Flament", "", "", "", "X", "X", ""),
    ("Persistance PostgreSQL, Redis, Celery, MinIO", "Gulraiz Mikail / Arthur Flament", "", "", "", "", "X", "X"),
    ("Hardening, recette, documentation et livraison", "Gulraiz Mikail + Arthur Flament", "", "", "", "", "", "X"),
]

PROJECT_ORG_ROWS = [
    ("Arthur Flament", "Referent interface, reporting et documentation", "46 j.h", "Pilotage du front, UX, artefacts PDF/JSON, recette et livrables de soutenance."),
    ("Gulraiz Mikail", "Referent backend, execution et securite", "50 j.h", "Services de scan, SSH Kali, orchestration, Docker, authentification et durcissement."),
    ("Commanditaire pedagogique", "Validation fonctionnelle et arbitrage", "8 j.h", "Cadre le besoin, valide les jalons et arbitre les priorites majeures."),
    ("Relecture technique", "Revue qualite et coherence du DAT", "6 j.h", "Controle la lisibilite du dossier, la coherence des preuves et la qualite de restitution."),
]

PROJECT_GOVERNANCE_ROWS = [
    ("Kick-off", "Semaine 1", "Equipe projet + commanditaire", "Validation du perimetre, des jalons, des risques et du mode operatoire."),
    ("Point projet", "Hebdomadaire", "Equipe projet", "Suivi avancement, blocages, reaffectation des taches et mise a jour du backlog."),
    ("Revue de sprint", "Toutes les 2 semaines", "Equipe projet + commanditaire", "Demonstration des increments livrables et validation des priorites suivantes."),
    ("COPIL de suivi", "Mensuel", "Equipe projet + commanditaire", "Arbitrage, verification budget/charge, suivi risques et decisions structurantes."),
    ("Recette finale", "Dernier mois", "Equipe projet + relecteur technique", "Validation documentaire, stabilisation de la demo et preparation de livraison."),
]

PROJECT_BACKLOG_ROWS = [
    ("Epics structurants", "12", "Architecture, execution, reporting, securite, persistance, documentation et demo."),
    ("User stories priorisees", "41", "Backlog cible organise par valeur demonstrative et faisabilite technique."),
    ("Taches techniques detaillees", "96", "Decoupage fin des travaux backend, frontend, infra, tests et livrables."),
    ("Jalons de validation", "4", "Kick-off, socle technique, beta fonctionnelle, livraison finale."),
]

PROJECT_BUDGET_ROWS = [
    ("Cadrage, architecture et conception", "14 j.h", "450 EUR / j.h", "6 300 EUR"),
    ("Developpement backend et orchestration", "32 j.h", "450 EUR / j.h", "14 400 EUR"),
    ("Frontend, reporting et experience operateur", "24 j.h", "420 EUR / j.h", "10 080 EUR"),
    ("Securite, tests et hardening", "16 j.h", "460 EUR / j.h", "7 360 EUR"),
    ("Documentation, recette et livraison", "12 j.h", "380 EUR / j.h", "4 560 EUR"),
    ("Infrastructure de lab et outillage", "Forfait", "-", "1 800 EUR"),
    ("Marge de pilotage et aleas", "10 % du total precedent", "-", "4 450 EUR"),
]

COMMAND_REFERENCE_ROWS = [
    ("Nmap", "Cible IP ou domaine", "nmap -Pn -sV --script http-title 192.168.159.1", "Commande representative lancee apres validation de la cible et de la allowlist."),
    ("Wireshark / tshark", "Interface + duree + cible", "tshark -n -i eth0 -a duration:10 -f host 192.168.159.145 -q -z io,stat,1 -z io,phs -z endpoints,ip", "Capture courte orientee synthese statistique et non conservation brute de PCAP."),
    ("WhatWeb", "URL ou service web", "whatweb http://192.168.159.1:3000", "Fingerprint simple de technologies visibles."),
    ("theHarvester", "Domaine uniquement", "theHarvester -d plutonia-mc.fr -b anubis,crtsh,rapiddns,urlscan -l 200", "Le service extrait ou valide le domaine avant execution."),
    ("Gobuster", "URL web de lab", "gobuster dir -u http://192.168.159.1:3000 -w /usr/share/wordlists/dirb/common.txt -t 4 --delay 150ms --timeout 5s -x php,html,js,txt", "Threads limites et delai ajoute pour proteger le lab."),
    ("SQLmap", "URL complete", "sqlmap -u https://cible.exemple/login?id=1 --batch --smart --random-agent --level 1 --risk 1 --threads 2", "Profil volontairement cadre pour limiter l'agressivite du scan."),
    ("Hydra SSH", "Hote + utilisateur(s) + mot(s) de passe", "hydra -s 22 -t 4 -L /usr/share/seclists/Usernames/top-usernames-shortlist.txt -P /usr/share/seclists/Passwords/Common-Credentials/10k-most-common.txt 192.168.159.145 ssh", "Le mot de passe manuel est masque dans l'affichage web."),
    ("Nikto", "URL ou hote web", "nikto -h http://192.168.159.1:3000", "Commande simple, potentiellement longue, avec timeout plus large."),
    ("SSLyze", "Service TLS", "sslyze 192.168.159.1:443", "Controle de posture TLS et de configuration cryptographique."),
    ("Auth Audit", "Page + endpoint de login", "auth-audit page=http://192.168.159.1:3000/#/login endpoint=http://192.168.159.1:3000/rest/user/login format=json tentatives=3", "Module natif base sur httpx, sans binaire externe."),
]

SECURITY_ENHANCEMENT_ROWS = [
    ("Secrets applicatifs", "Variables d'environnement en clair pour plusieurs secrets", "Basculer vers Fernet ou secret store, rotation des secrets, separation des secrets par environnement"),
    ("Transport HTTP", "Usage local acceptable sans HTTPS fort", "Activer TLS reverse proxy, cookies secure, HSTS et durcissement de session"),
    ("Controle d'acces", "Authentification locale + TOTP, pas de RBAC fin", "Ajouter RBAC par role : administration, operateur, lecture seule"),
    ("Identite et MFA", "Compte local unique avec activation TOTP stockee en local", "Passer a des comptes persistants en base, enrollment MFA par utilisateur, codes de secours et revocation admin"),
    ("Journalisation", "Audit de base sur les actions critiques", "Structurer les audit logs, retention, correlation et export SIEM"),
    ("Execution distante", "Kali joignable par SSH avec confiance configurable", "Gestion stricte des host keys, comptes dedies, bastion ou reseau d'administration isole"),
    ("Commandes", "Validation forte sur plusieurs modules mais pas sandbox universelle", "Generaliser les allowlists, quotas et profils d'execution par module"),
]

ERGONOMIC_ENHANCEMENT_ROWS = [
    ("Interface principale", "Conserver une page sobre mais ajouter une aide contextuelle par module"),
    ("Historique", "Ajouter filtres multicriteres, tags, favoris et comparaison entre sessions"),
    ("Console", "Ajouter repli par module, surlignage stdout/stderr et export direct de traces"),
    ("Rapports", "Ajouter une synthese executive, une vue technique detaillee et une table des matieres interne"),
    ("Rapports enrichis", "Introduire un pipeline de synthese assistee avec validation humaine avant export final"),
    ("Campagnes", "Permettre la sauvegarde de profils de scans et la relance d'une session type"),
    ("Accessibilite", "Mieux gerer contrastes, navigation clavier et lisibilite des formulaires avances"),
]


TOOLS = [
    {
        "number": "10.2.1",
        "title": "Nmap",
        "command": "nmap -Pn -sV [--script ...] <target>",
        "pole": "Sécurité (SOC, EDR, XDR)",
        "body": dedent(
            """
            Nmap constitue le socle de reconnaissance réseau de la toolbox. Son intégration répond au besoin initial de proposer, dès la première itération, un outil robuste et immédiatement démontrable sur des cibles d'infrastructure, des postes de test ou des services exposés au sein d'un lab. Le module a été volontairement cadré autour d'un noyau d'arguments fixes, à savoir l'option -Pn pour éviter la dépendance au ping préalable et l'option -sV pour enrichir le résultat par une détection de versions de services.

            L'enjeu principal n'était pas seulement d'appeler la commande nmap, mais de la faire entrer dans un cadre d'exécution sécurisé. À ce titre, l'application n'expose pas un champ libre permettant de transmettre des arguments arbitraires à l'outil. Les scripts NSE autorisés sont décrits dans le catalogue du module et validés par une allowlist stricte. Cette approche réduit le risque d'usage non maîtrisé, normalise les démonstrations et simplifie l'interprétation des rapports.

            Sur le plan applicatif, Nmap repose sur la classe commune de streaming. L'utilisateur coche le module dans l'interface, saisit une cible, puis suit l'exécution ligne par ligne dans le terminal web. À l'issue du scan, les données sont transformées en artefact JSON et en PDF de synthèse. L'événement d'exécution est également journalisé dans le journal d'audit, ce qui permet de conserver la trace du module invoqué, de la cible et du code de retour.

            L'intégration de Nmap reflète bien l'esprit général de Reconforge : rendre des outils puissants plus accessibles, sans perdre le contrôle sur leur exécution. Le module a aussi servi de base de travail pour le reste de la plateforme, notamment sur la validation des entrées, le streaming temps réel, la persistance des résultats et l'agrégation dans une session unique lorsque plusieurs outils sont enchaînés.
            """
        ).strip(),
        "capture": "[Capture - Exécution Nmap avec sortie temps réel dans la console]",
    },
    {
        "number": "10.2.2",
        "title": "Wireshark / tshark",
        "command": "tshark -n -i <interface> -a duration:<n> -f host <target> -q -z ...",
        "pole": "Sécurité (SOC, EDR, XDR)",
        "body": dedent(
            """
            Le module Wireshark a été implémenté non pas sous la forme d'une capture graphique classique, mais via tshark, son équivalent en ligne de commande. Ce choix est cohérent avec une architecture web containerisée : l'objectif n'est pas d'ouvrir une interface locale sur le poste de l'utilisateur, mais de produire un résumé exploitable dans un rapport et dans un terminal web. L'intégration vise principalement des écoutes courtes de quelques secondes, orientées démonstration ou vérification d'un flux.

            La cible saisie dans l'interface est réutilisée comme filtre réseau de type host. L'utilisateur renseigne également l'interface réseau à écouter et la durée, en secondes, de la capture. Afin de limiter les dérives, la durée a été bornée entre une seconde et cinq minutes. Le service ajuste par ailleurs son timeout interne pour laisser un léger tampon au-dessus de la durée demandée, ce qui évite qu'une capture valide ne soit coupée prématurément par l'application elle-même.

            Les données retournées ne correspondent pas à un fichier PCAP brut. Le module restitue plutôt des statistiques d'entrée/sortie, une hiérarchie de protocoles et un résumé des endpoints IP observés. Cette décision a deux avantages : elle réduit le volume de données à conserver et elle simplifie la lecture du résultat côté utilisateur. En contrepartie, la toolbox n'a pas vocation, à ce stade, à remplacer un poste d'analyse forensique complet.

            Ce module rappelle que la plateforme ne se limite pas aux seuls scans web. Elle peut aussi couvrir une partie de l'observation réseau, à condition de le faire dans un format court, lisible et compatible avec une interface de démonstration locale.
            """
        ).strip(),
        "capture": "[Capture - Paramétrage Wireshark avec durée manuelle et exemple de rapport]",
    },
    {
        "number": "10.3.1",
        "title": "WhatWeb",
        "command": "whatweb <target>",
        "pole": "Développement SaaS",
        "body": dedent(
            """
            WhatWeb a été retenu comme outil de fingerprint web rapide, particulièrement utile pour qualifier une surface applicative avant de lancer des contrôles plus lourds. Dans la toolbox, il sert de module d'observation : technologies de front-end détectées, serveurs identifiés, composants visibles et indices de stack. Ce type de résultat oriente ensuite l'usage d'outils plus spécialisés, comme Nikto ou SQLmap.

            L'intégration du module est volontairement sobre. Lorsqu'une URL complète est fournie, le service conserve l'URL telle quelle ; lorsqu'un domaine ou une IP est saisi, la cible est normalisée avant exécution. Aucun paramétrage avancé n'est exposé dans l'interface, ce qui correspond à l'objectif du projet : fournir une entrée simple, un outil identifiable et un retour immédiat. Cette simplicité améliore la démonstration et évite de surcharger la page principale.

            Techniquement, WhatWeb hérite de la classe d'exécution commune. Son intérêt dans l'architecture du backend est donc double : il apporte un cas d'usage concret côté reconnaissance web et il confirme que le patron de développement retenu reste stable d'un module à l'autre. Le coût d'intégration d'un nouveau binaire demeure limité dès lors que le service sait normaliser la cible, construire la commande et produire les messages d'erreur adaptés.

            En phase de présentation, WhatWeb permet également de matérialiser rapidement la valeur de la console temps réel. L'utilisateur voit la commande partir, puis obtient une sortie concise, souvent plus lisible que celle de scanners lourds. Le module joue ainsi un rôle pédagogique fort dans la compréhension du fonctionnement de la toolbox.
            """
        ).strip(),
        "capture": "[Capture - Module WhatWeb sélectionné et résultat de fingerprint]",
    },
    {
        "number": "10.3.2",
        "title": "theHarvester",
        "command": "theHarvester -d <domaine> -b anubis,crtsh,rapiddns,urlscan -l 200",
        "pole": "Développement SaaS",
        "body": dedent(
            """
            theHarvester vient enrichir le pôle Développement SaaS en apportant une dimension OSINT. Contrairement aux modules qui interrogent directement un service HTTP en ligne, celui-ci travaille principalement à partir d'un domaine et de sources publiques, comme des bases de transparence de certificats, des moteurs d'indexation ou des annuaires techniques. L'outil a donc été classé dans une logique de préparation et de cartographie d'exposition externe.

            L'application extrait automatiquement le nom de domaine lorsqu'une URL complète est saisie. Ce comportement a été ajouté pour éviter à l'utilisateur d'avoir à reformuler sa cible manuellement. À l'inverse, le service refuse explicitement les adresses IP pures, car elles ne correspondent pas au mode de fonctionnement normal de theHarvester. Cette validation métier participe à la qualité du résultat et évite des exécutions inutiles.

            Le choix des sources par défaut a été volontairement restreint à un ensemble pertinent et relativement stable : anubis, crtsh, rapiddns et urlscan. Une limite de résultats a également été fixée afin de garder des temps de réponse compatibles avec une démonstration pédagogique. Là encore, l'objectif n'est pas d'exposer toute la richesse native de l'outil dès la première version, mais de fournir un module exploitable dans le cadre d'une session multi-outils.

            En termes d'architecture, l'intérêt de theHarvester est de démontrer que la toolbox sait intégrer des outils de natures variées. Tous les modules ne se comportent pas comme des scanners actifs ; certains réalisent de la collecte d'information indirecte. Le backend a été conçu de manière suffisamment générique pour absorber cette diversité sans remettre en cause la structure applicative.
            """
        ).strip(),
        "capture": "[Capture - theHarvester dans la catégorie Développement SaaS]",
    },
    {
        "number": "10.3.3",
        "title": "Gobuster",
        "command": "gobuster dir -u <url> -w <wordlist> -t 4 --delay 150ms --timeout 5s ...",
        "pole": "Développement SaaS",
        "body": dedent(
            """
            Gobuster a été ajouté pour couvrir les besoins d'énumération de contenu web, tout en restant compatible avec un laboratoire de démonstration fragile comme Juice Shop. Le risque principal rencontré lors des essais était la saturation de la cible et, par ricochet, le blocage de l'interface de la toolbox. Pour cette raison, le module a été volontairement bridé par défaut : nombre de threads réduit, délai entre les requêtes et timeout limité.

            Une logique complémentaire de détection de réponses génériques a également été intégrée. Avant d'exécuter l'énumération, le service interroge une URL aléatoire afin d'observer le code HTTP et la taille d'une réponse à une ressource inexistante. Si la cible renvoie tout de même une page valide, la longueur détectée est transmise à Gobuster via l'option d'exclusion appropriée. Cette adaptation limite les faux positifs sur des applications qui répondent 200 à des chemins arbitraires.

            Du point de vue de l'expérience utilisateur, le module ne demande aucun réglage supplémentaire dans l'interface. Cela correspond à l'arbitrage fonctionnel retenu : exposer un usage standard, maîtrisé et démontrable, plutôt qu'un configurateur complet de bruteforce HTTP. La complexité métier reste dans le backend, là où les garde-fous peuvent être centralisés et maintenus.

            Gobuster occupe une place importante dans le DAT car il illustre bien la manière dont le projet traite les contraintes de performance et de stabilité. L'intégration n'a pas consisté à encapsuler naïvement un binaire, mais à adapter ses paramètres pour qu'il reste supportable dans un environnement local, pédagogique et conteneurisé.
            """
        ).strip(),
        "capture": "[Capture - Exécution Gobuster avec exclusion automatique de wildcard]",
    },
    {
        "number": "10.3.4",
        "title": "SQLmap",
        "command": "sqlmap -u <url> --batch --smart --random-agent --level 1 --risk 1 --threads 2",
        "pole": "Développement SaaS",
        "body": dedent(
            """
            SQLmap couvre le besoin de test applicatif orienté injections SQL. Son intégration a toutefois été volontairement simplifiée afin d'éviter deux écueils : d'une part une interface trop lourde, d'autre part des campagnes de test trop agressives pour une cible de laboratoire. Le module exige donc une URL complète et exécute un profil standard de commandes, non interactif, avec des niveaux de risque et de profondeur modérés.

            Le choix des options retenues n'est pas anodin. Le mode --batch supprime les interactions bloquantes, --smart réduit certains tests redondants, --random-agent améliore la compatibilité face à quelques protections simples, et les paramètres --level 1, --risk 1 et --threads 2 limitent l'impact de la commande. Cette combinaison produit un comportement acceptable pour une démonstration sans prétendre couvrir l'ensemble des capacités de SQLmap.

            L'expérience de test a montré que l'outil est particulièrement pertinent lorsqu'il cible des URL porteuses de paramètres explicites, ou des applications vulnérables prévues à cet effet. À l'inverse, sur des pages de connexion complexes ou fortement filtrées, le résultat peut légitimement conclure à l'absence d'injectabilité exploitable. Le DAT doit mettre ce point en avant afin de replacer l'outil dans son bon contexte d'utilisation.

            D'un point de vue architectural, SQLmap confirme que la toolbox sait aussi intégrer des outils à sortie volumineuse. Le flux temps réel, l'agrégation dans une session unique et la conservation de la sortie standard dans les rapports permettent à l'utilisateur de conserver la trace complète d'une analyse sans quitter l'interface.
            """
        ).strip(),
        "capture": "[Capture - Session SQLmap et rapport agrégé généré par la toolbox]",
    },
    {
        "number": "10.4.1",
        "title": "Hydra (mode SSH)",
        "command": "hydra -s <port> -t 4 -l|-L ... -p|-P ... <target> ssh",
        "pole": "Infrastructure",
        "body": dedent(
            """
            Hydra a été retenu pour représenter les tests d'authentification de type infrastructure. Dans la version actuelle, le périmètre a été volontairement limité au protocole SSH. Cette décision résulte directement des expérimentations menées pendant le projet : l'usage de Hydra sur des formulaires web modernes s'est révélé instable, lent et peu lisible dans le contexte de démonstration retenu. Le cadrage SSH permet en revanche un comportement plus prévisible et plus simple à expliquer.

            L'interface expose uniquement les paramètres nécessaires à ce mode ciblé : identifiant ou source d'identifiants, mot de passe ou wordlist, ainsi que le port SSH. Les listes prédéfinies s'appuient sur des fichiers Kali courants, comme top-usernames-shortlist.txt, names.txt, 500-worst-passwords.txt, 10k-most-common.txt ou rockyou.txt. Cette approche répond à un besoin pédagogique clair : montrer comment la toolbox peut piloter un outil de test d'authentification tout en laissant un choix simple à l'utilisateur.

            Le backend masque le mot de passe manuel dans la commande affichée, afin d'éviter la fuite visuelle d'un secret dans la console. Il vérifie également la cohérence des modes sélectionnés et refuse les combinaisons incomplètes ou incohérentes. Le module ne cherche donc pas à tout permettre ; il cherche à fournir une expérience robuste, contrôlée et adaptée à un environnement de lab.

            Ce module illustre bien la manière dont Reconforge arbitre entre richesse fonctionnelle et maîtrise opérationnelle. Hydra est bien exposé dans l'interface, mais dans une forme volontairement contrainte, plus réaliste pour un usage local et plus soutenable pour la démonstration.
            """
        ).strip(),
        "capture": "[Capture - Panneau Hydra affiché uniquement lorsque la case est cochée]",
    },
    {
        "number": "10.5.1",
        "title": "Nikto",
        "command": "nikto -h <target>",
        "pole": "Support client",
        "body": dedent(
            """
            Nikto a été intégré afin de couvrir un besoin d'audit rapide sur des applications ou serveurs web exposés. Son positionnement dans le pôle Support client s'explique par son intérêt pratique pour des vérifications de configuration, de contenus accessibles et de signatures connues sur des outils de communication ou des portails internes. L'objectif n'est pas de concurrencer un scanner applicatif complet, mais de proposer un contrôle reconnaissable, utile et immédiatement lisible.

            L'intégration backend suit le même schéma que les autres modules de scan actif. La cible est normalisée, la commande est construite, puis la sortie est streamée en direct vers le terminal web. L'utilisateur n'a pas à connaître la syntaxe de Nikto pour lancer un premier audit : il se concentre sur la cible et sur l'interprétation du résultat. Ce niveau d'abstraction constitue un argument fort pour la valeur pédagogique de la toolbox.

            Les essais ont montré que Nikto peut être plus long que d'autres modules. Ce comportement est cohérent avec la nature de l'outil, qui multiplie les vérifications et les signatures. Le projet a donc retenu un timeout plus large pour ce service. Ce point est explicité dans le DAT, car il montre que les modules ne sont pas traités de façon uniforme ; chacun dispose d'un cadrage spécifique selon sa charge et son profil d'exécution.

            Enfin, Nikto participe à l'intérêt de l'historique et des rapports. La verbosité de sa sortie rend utile la conservation d'un artefact consultable après exécution, ce qui justifie pleinement le mécanisme de persistance retenu dans la plateforme.
            """
        ).strip(),
        "capture": "[Capture - Exécution Nikto et sortie détaillée dans l'historique]",
    },
    {
        "number": "10.5.2",
        "title": "SSLyze",
        "command": "sslyze <target>",
        "pole": "Support client",
        "body": dedent(
            """
            SSLyze répond à un besoin simple mais fréquent : vérifier la posture TLS d'un service exposé. Son intégration permet de tester rapidement des configurations de certificats, des suites cryptographiques et des paramètres de négociation sans sortir de la toolbox. Dans un projet à visée démonstrative, ce module a l'avantage d'apporter un résultat compréhensible, souvent plus ciblé qu'un scan réseau généraliste.

            Le choix de l'intégrer dans le pôle Support client s'appuie sur le fait que de nombreux outils internes ou portails utilisateurs reposent sur HTTPS. La qualité de cette couche de transport est donc un sujet concret, directement observable et facile à justifier dans un DAT. L'application n'expose pas d'options avancées complexes ; elle fournit un point d'entrée simple vers un diagnostic TLS de premier niveau.

            Comme pour les autres modules standards, SSLyze s'insère dans la chaîne d'exécution commune. Cela permet de mutualiser la gestion du streaming, du timeout, du reporting et de la journalisation. La répétition de ce patron architectural n'est pas un défaut ; au contraire, elle montre que la plateforme repose sur un cadre cohérent et maintenable.

            Sur le plan projet, SSLyze aide à démontrer que la toolbox ne se limite pas à la recherche de vulnérabilités applicatives. Elle couvre aussi des contrôles de configuration et de robustesse technique, ce qui renforce sa crédibilité comme base de travail orientée pentest et audit.
            """
        ).strip(),
        "capture": "[Capture - Résultat SSLyze avec rapport PDF associé]",
    },
    {
        "number": "10.6.1",
        "title": "Auth Audit",
        "command": "auth-audit page=<url> endpoint=<url> format=<json|form> tentatives=<n>",
        "pole": "RH / Administration",
        "body": dedent(
            """
            Auth Audit est un module spécifique développé pour le projet afin de traiter le besoin d'évaluation contrôlée d'un parcours de connexion web, sans basculer vers un bruteforce agressif. Son rôle n'est pas de casser un mot de passe, mais d'observer la manière dont un formulaire ou un endpoint d'authentification réagit à quelques tentatives invalides : codes de retour, homogénéité des messages, présence de cookies, détection de CAPTCHA ou de MFA, et indices de limitation de débit.

            L'intégration de ce module répond à un besoin particulièrement pertinent pour le pôle RH / Administration, dans lequel de nombreux outils internes exposent des interfaces de gestion sensibles. Le paramétrage de l'interface permet de renseigner l'URL de la page de connexion, l'endpoint d'authentification, le format de la requête, les noms des champs identifiant et mot de passe, un marqueur d'échec et le nombre de tentatives contrôlées. Cette granularité rend le module adaptable à différents cas, sans exiger de l'utilisateur qu'il écrive lui-même une logique de test.

            Techniquement, Auth Audit se distingue des autres modules car il n'appelle pas un binaire externe. Il exécute directement des requêtes HTTP via httpx, collecte les réponses et produit un résumé interprétable. Cela démontre que la toolbox n'est pas limitée aux wrappers de CLI. Elle peut également embarquer des modules natifs, spécialement développés pour un besoin précis, tout en conservant le même contrat de sortie que les autres services.

            Ce module incarne aussi une approche responsable du test. Dans un cadre académique ou de démonstration, il est souvent plus pertinent d'évaluer un parcours d'authentification par son comportement que de chercher à lancer une force brute lourde, lente et difficile à soutenir sur des cibles web modernes.
            """
        ).strip(),
        "capture": "[Capture - Formulaire Auth Audit et résumé d'audit d'authentification]",
    },
]


def configure_section(section):
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def ensure_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = tc_mar.find(qn(f"w:{tag}"))
        if item is None:
            item = OxmlElement(f"w:{tag}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BORDER)


def set_table_width(table, total_width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_column_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_page_field(paragraph):
    run = paragraph.add_run("Page ")
    ensure_font(run, size=9, color=MUTED)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def create_styles(document: Document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    if "DAT Metadata" not in styles:
        style = styles.add_style("DAT Metadata", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10)
        style.font.color.rgb = MUTED
        style.paragraph_format.space_after = Pt(3)

    if "DAT Placeholder" not in styles:
        style = styles.add_style("DAT Placeholder", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10)
        style.font.color.rgb = DARK_BLUE
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)


def add_header_footer(section, first_page=False):
    if first_page:
        section.different_first_page_header_footer = True
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"DAT - {APP_NAME}")
    ensure_font(run, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)


def add_paragraphs(document: Document, text: str):
    for block in [item.strip() for item in text.split("\n\n") if item.strip()]:
        p = document.add_paragraph()
        run = p.add_run(block)
        ensure_font(run)


def add_bullet_list(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        run = p.add_run(item)
        ensure_font(run)


def add_placeholder(document: Document, text: str):
    p = document.add_paragraph(style="DAT Placeholder")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    ensure_font(run, size=10, color=DARK_BLUE, italic=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), PLACEHOLDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_table(document: Document, headers: list[str], rows: list[tuple[str, ...]], widths_cm: list[float]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table)
    set_column_widths(table, widths_cm)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        cell = header_cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        ensure_font(run, size=10, bold=True)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            ensure_font(run, size=10)

    document.add_paragraph()
    return table


def add_cover(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    run = p.add_run("DOCUMENT D'ARCHITECTURE TECHNIQUE")
    ensure_font(run, name="Calibri", size=14, color=MUTED, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(APP_NAME)
    ensure_font(run, name="Calibri", size=28, color=DARK_BLUE, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(APP_TAGLINE)
    ensure_font(run, size=14, color=BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(
        "Ce document décrit l'architecture technique de la solution réalisée, les choix de conception retenus, "
        "leur justification au regard du cahier des charges, ainsi que les limites et perspectives d'évolution de la plateforme."
    )
    ensure_font(run, size=11)

    table = document.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table)
    set_column_widths(table, [4.0, 11.5])
    metadata = [
        ("Projet", APP_NAME),
        ("Type de document", "DAT - version rédigée"),
        ("Version", "1.0"),
        ("Date", datetime.now().strftime("%d/%m/%Y")),
        ("Périmètre", "Architecture, sécurité, exécution des outils, reporting, exploitation et trajectoire d'industrialisation"),
        ("Statut", "Document d'architecture technique pour présentation et soutenance"),
    ]
    for row_idx, (label, value) in enumerate(metadata):
        left, right = table.rows[row_idx].cells
        for cell in (left, right):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, LIGHT_FILL)
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        lrun = lp.add_run(label)
        rrun = rp.add_run(value)
        ensure_font(lrun, size=10, bold=True)
        ensure_font(rrun, size=10)

    document.add_paragraph()
    add_placeholder(document, "[Capture - Page d'accueil de la toolbox ou visuel de couverture du projet]")
    document.add_page_break()


def add_document_control(document: Document):
    document.add_paragraph("Contrôle documentaire", style="Heading 1")
    add_paragraphs(
        document,
        dedent(
            """
            Le présent DAT a été rédigé pour servir de support de présentation, de justification technique et de base de maintenance pour la Pentest Toolbox. Il ne s'agit ni d'une note d'intention, ni d'un simple plan de projet, mais bien d'un document technique décrivant une solution réellement implémentée dans un environnement local conteneurisé.

            Le document distingue volontairement ce qui est déjà opérationnel, ce qui est préparé dans l'architecture mais pas encore activé dans le flux courant, et ce qui relève de la feuille de route. Cette distinction est essentielle pour conserver une lecture honnête du projet et pour rendre la démonstration défendable face à un évaluateur ou à un commanditaire.
            """
        ).strip(),
    )

    add_table(
        document,
        ["Version", "Date", "Auteur", "Nature de la modification"],
        [("1.0", datetime.now().strftime("%d/%m/%Y"), "Arthur / Codex", "Création de la version rédigée complète du DAT")],
        [2.2, 3.0, 4.0, 7.3],
    )

    document.add_paragraph("Résumé exécutif", style="Heading 1")
    add_paragraphs(
        document,
        dedent(
            """
            La Pentest Toolbox est une application web locale conçue pour centraliser l'exécution d'outils de pentest au sein d'une interface unique, lisible et pilotable. L'utilisateur saisit une cible, sélectionne un ou plusieurs modules puis suit l'exécution des commandes en temps réel dans un terminal web. Les sorties sont consolidées dans une session unique, historisées et exportées sous forme de rapports JSON et PDF.

            L'architecture repose sur Python 3.12, FastAPI, Jinja2 et Docker Compose. Le socle d'infrastructure inclut PostgreSQL, Redis et MinIO, afin d'aligner la solution avec le cadrage technique demandé, même si la persistance active des rapports reste aujourd'hui fondée sur le stockage local dans le répertoire storage. Ce choix intermédiaire permet de livrer une solution fonctionnelle rapidement tout en préparant une trajectoire d'industrialisation crédible.

            Un point structurant du projet est la séparation entre l'interface web et l'environnement d'exécution des outils. La toolbox peut fonctionner localement, mais elle a surtout été adaptée pour déléguer les commandes à une machine Kali Linux joignable en SSH. Cette décision répond à une contrainte pratique forte : de nombreux outils de sécurité sont déjà disponibles et mieux maintenus dans Kali que dans un conteneur web générique.

            Sur le volet sécurité, l'application dispose d'une authentification locale optionnelle avec mot de passe et double facteur TOTP. Un écran de configuration initiale permet d'activer l'A2F via QR code. Les scans exécutés sont journalisés dans un audit log dédié. En revanche, certaines briques du cahier des charges ne sont aujourd'hui que préparées et non complètement activées, notamment RBAC, HTTPS généralisé et chiffrement Fernet des secrets. Le DAT expose clairement ces écarts pour préserver la cohérence du discours technique.
            """
        ).strip(),
    )
    document.add_page_break()


def add_manual_toc(document: Document):
    document.add_paragraph("Sommaire de lecture", style="Heading 1")
    add_paragraphs(
        document,
        "Le sommaire ci-dessous présente l'organisation logique du document. La version finale pourra être enrichie d'une table des matières Word mise à jour avant remise définitive."
    )
    toc_items = [
        "1. Introduction et objectifs du document",
        "2. Contexte projet et rappel du cahier des charges",
        "3. Réponse au cahier des charges et niveau de couverture",
        "4. Principes directeurs de l'architecture",
        "5. Architecture d'infrastructure et conteneurisation",
        "6. Architecture applicative et logique backend/frontend",
        "7. Mécanisme d'exécution des outils et délégation Kali",
        "8. Sécurité de la plateforme et authentification forte",
        "9. Stockage, reporting, historisation et persistance",
        "10. Description détaillée des modules par pôle",
        "11. Maintenabilité, extensibilité et industrialisation",
        "12. Exploitation, déploiement et bonnes pratiques d'usage",
        "13. Scénarios de test, validations et limites observées",
        "14. Risques, écarts actuels et feuille de route",
        "15. Conclusion",
        "16. Annexes techniques",
    ]
    add_bullet_list(document, toc_items)
    document.add_page_break()


def add_heading(document: Document, title: str, level: int = 1):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    document.add_paragraph(title, style=style)


def add_requirement_matrix(document: Document):
    add_heading(document, "3. Réponse structurée au cahier des charges", 1)
    add_paragraphs(
        document,
        dedent(
            """
            La conformité à un cahier des charges technique ne se réduit pas à une liste de technologies installées. Dans le cadre de ce projet, la réponse a été construite autour d'un principe de cohérence : chaque choix retenu devait soit répondre directement à une exigence exprimée, soit préparer une exigence encore non industrialisée sans compromettre la démonstration de la plateforme.

            La matrice suivante distingue les exigences totalement couvertes, les exigences partiellement couvertes et les exigences encore en préparation. Ce niveau de transparence est indispensable, car il permet de présenter un projet réaliste : la solution est fonctionnelle, démontrable et évolutive, mais elle n'est pas artificiellement présentée comme achevée sur l'ensemble des ambitions initiales.
            """
        ).strip(),
    )
    add_table(
        document,
        ["Domaine", "Exigence", "Statut", "Réponse apportée"],
        REQUIREMENT_ROWS,
        [3.0, 5.2, 2.0, 6.3],
    )
    add_paragraphs(
        document,
        dedent(
            """
            Les exigences déjà couvertes concernent surtout le socle applicatif : Python 3.x, framework web moderne, architecture modulaire, documentation Swagger, conteneurisation via Docker Compose et intégration progressive d'outils opérationnels. Les exigences partiellement couvertes sont liées à la maturation de la solution : persistance relationnelle effective, stockage objet réellement branché, orchestration asynchrone des scans, contrôle d'accès par rôles et durcissement complet du volet sécurité.

            Cette situation est cohérente avec la stratégie de livraison retenue. Plutôt que de produire une architecture théorique difficile à démontrer, le projet a privilégié une base fonctionnelle, puis a préparé les briques d'extension dans l'environnement Docker et dans la structure logicielle. Le DAT doit être lu dans cette perspective : il documente une plateforme livrée en capacité opérationnelle initiale, sur laquelle un cycle d'industrialisation est clairement envisageable.
            """
        ).strip(),
    )


def add_tool_sections(document: Document):
    add_heading(document, "10. Description détaillée des modules intégrés", 1)
    add_paragraphs(
        document,
        dedent(
            """
            L'intégration des outils constitue le cœur visible de la toolbox. Chaque module a été ajouté avec le même objectif général, à savoir encapsuler un outil de pentest ou d'audit dans une interface plus simple à manipuler, tout en conservant la traçabilité des commandes exécutées et la lisibilité des résultats. Le projet n'a pas cherché à exposer la totalité des options natives de chaque outil ; il a cherché à proposer une version cohérente, pédagogique et suffisamment robuste pour une démonstration concrète.

            Le classement des modules par pôles répond à un double enjeu. D'une part, il reprend la logique métier évoquée dans le cadrage initial, ce qui facilite la lecture du projet par un évaluateur non développeur. D'autre part, il structure l'interface et évite une simple liste plate d'outils. Cette organisation permet de faire le lien entre le besoin utilisateur, l'usage de l'application et la manière dont le code est lui-même organisé dans le registre des modules.
            """
        ).strip(),
    )
    add_table(
        document,
        ["Pôle", "Module", "Rôle dans la plateforme"],
        MODULE_ROWS,
        [4.1, 3.3, 9.1],
    )
    add_placeholder(document, "[Capture - Vue des modules classés par pôle dans l'interface]")

    for tool in TOOLS:
        add_heading(document, f"{tool['number']} {tool['title']}", 2)
        p = document.add_paragraph()
        lead = p.add_run("Pôle concerné : ")
        ensure_font(lead, bold=True)
        body = p.add_run(tool["pole"])
        ensure_font(body)

        p = document.add_paragraph()
        lead = p.add_run("Commande type : ")
        ensure_font(lead, bold=True)
        body = p.add_run(tool["command"])
        ensure_font(body)

        add_paragraphs(document, tool["body"])
        add_placeholder(document, tool["capture"])


def build_document():
    document = Document()
    configure_section(document.sections[0])
    create_styles(document)
    add_header_footer(document.sections[0], first_page=True)

    add_cover(document)
    add_document_control(document)
    add_manual_toc(document)

    add_heading(document, "1. Introduction", 1)
    add_heading(document, "1.1 Contexte du projet", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La Pentest Toolbox est née d'un besoin opérationnel simple : disposer d'une interface locale capable de centraliser plusieurs outils de sécurité, de normaliser leur exécution et d'en restituer les résultats dans un format exploitable. Dans de nombreux contextes pédagogiques ou de démonstration, l'utilisation directe de la ligne de commande constitue un frein. Les outils sont puissants, mais ils supposent de maîtriser leur syntaxe, leur environnement d'installation et parfois leurs dépendances spécifiques. Le projet vise donc à réduire cette friction sans masquer la réalité technique de ce qui est exécuté.

            La plateforme a été pensée pour fonctionner en localhost, dans un environnement Docker, avec une interface volontairement sobre : un champ unique pour la cible, des cases à cocher pour les outils, une console temps réel et un historique des sessions. Cette sobriété n'est pas un choix esthétique secondaire ; elle répond à un objectif d'ergonomie. L'utilisateur doit accéder rapidement à l'essentiel sans être noyé dans une interface de scanner professionnel difficile à défendre dans le cadre d'un livrable académique.

            Le projet s'inscrit également dans une logique de présentation propre et crédible. Le dossier devait être suffisamment propre pour être montré, ce qui a influencé plusieurs décisions techniques : nommage explicite des scripts, structure claire du backend, séparation nette entre routes, services et modules, et limitation volontaire des options exposées dans la page principale. Le DAT présenté ici formalise précisément ces arbitrages.
            """
        ).strip(),
    )
    add_placeholder(document, "[Capture - Interface principale de la toolbox au chargement]")

    add_heading(document, "1.2 Objectifs du document", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le présent Document d'Architecture Technique a pour objectif de décrire la solution telle qu'elle a été mise en œuvre. Il ne s'agit pas d'un cahier d'idées ni d'une projection abstraite. Les composants décrits dans ce document correspondent aux éléments réellement codés, configurés et testés dans l'environnement du projet, sauf mention explicite d'une capacité encore préparée mais non exploitée dans le flux courant.

            Le DAT répond à quatre fonctions principales. Premièrement, il sert de support de justification vis-à-vis du cahier des charges initial. Deuxièmement, il sert de dossier de compréhension technique pour une personne qui reprendrait le projet. Troisièmement, il sert de base de discussion pour les évolutions futures, notamment celles liées à la persistance, à l'orchestration asynchrone et au durcissement de la sécurité. Enfin, il sert d'appui de démonstration lors d'une soutenance ou d'une présentation.

            Pour remplir ces objectifs, le document adopte un ton volontairement professionnel et précis. Les choix sont justifiés, les limites sont exposées et les écarts entre cible d'architecture et état courant sont explicitement signalés. Cette transparence est essentielle pour donner au projet une crédibilité technique réelle.
            """
        ).strip(),
    )

    add_heading(document, "1.3 Périmètre couvert par la solution", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le périmètre actuellement couvert par la toolbox comprend l'interface web locale, l'API FastAPI, le streaming temps réel des sorties d'outils, la génération de rapports de session, l'historique des exécutions, la délégation des commandes vers une machine Kali Linux et une authentification locale optionnelle renforcée par un second facteur TOTP. Sur le plan des modules, dix outils sont intégrés à ce jour, répartis par pôles fonctionnels.

            En revanche, le périmètre n'inclut pas encore certaines briques pourtant mentionnées dans le cadrage cible. C'est le cas d'un RBAC complet, d'un chiffrement applicatif effectif des secrets via Fernet, d'un report d'artefacts dans MinIO, d'une persistance relationnelle active des scans dans PostgreSQL, d'une exécution réellement asynchrone via Celery pour les scans utilisateurs et d'une publication sécurisée en HTTPS. Ces sujets figurent donc à la fois dans l'analyse de couverture et dans la feuille de route.

            Cette délimitation est importante, car elle évite deux contresens fréquents : confondre présence d'un conteneur avec usage effectif dans le flux métier, et confondre modularité prête à l'emploi avec fonctionnalité déjà finalisée. Le DAT pose un cadre clair pour éviter ces ambiguïtés.
            """
        ).strip(),
    )

    add_heading(document, "2. Contexte projet et rappel du cahier des charges", 1)
    add_heading(document, "2.1 Attentes fonctionnelles initiales", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le besoin initial exprimé reposait sur la création d'une interface web en localhost permettant d'exécuter des scripts de sécurité à partir d'une cible unique saisie par l'utilisateur. L'interface devait rester simple : un champ principal pour une IP, un domaine ou une URL, puis une zone de sélection des scripts à lancer. Cette simplicité d'entrée a été maintenue jusqu'au bout du projet, car elle constitue l'un des éléments distinctifs de la solution.

            Une autre attente forte concernait la visibilité de l'exécution. Il ne s'agissait pas seulement de déclencher un scan, mais de montrer en direct la commande exécutée et ce qu'elle produit. Cette exigence a orienté très tôt le choix du streaming des sorties standard et d'erreur, puis la construction d'une console web. Elle a également conduit à la notion de session multi-outils, afin qu'une campagne de test apparaisse comme un ensemble cohérent et non comme une succession désordonnée d'exécutions isolées.

            Enfin, le besoin incluait un historique lisible des commandes déjà lancées, avec un accès aux résultats exportés. Cette demande a été traduite en une zone d'historique compacte, dotée d'une recherche simple, et reliée à un stockage local des rapports. Le projet s'est donc structuré autour d'une promesse claire : centraliser, rendre lisible, conserver et présenter.
            """
        ).strip(),
    )

    add_heading(document, "2.2 Exigences techniques conseillées", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le cadrage technique recommandait une architecture modulaire et sécurisée fondée sur Python 3.x, Flask ou FastAPI, Poetry, PostgreSQL, MinIO, Redis, Celery, Jinja2, Docker Compose et une API documentée. Il recommandait également l'intégration progressive d'outils tels que Nmap, OpenVAS, Nessus, Metasploit, ZAP, SQLmap, Hydra ou Aircrack-ng, avec un niveau d'industrialisation permettant une évolution vers le reporting et la gestion de rôles.

            Le projet a choisi de s'aligner largement sur cette recommandation, avec une nuance importante : l'architecture a d'abord été réalisée autour d'un noyau fonctionnel concret, puis complétée par des briques prêtes à l'extension. Cette méthode a évité de créer une stack techniquement complète mais peu démontrable. Par exemple, PostgreSQL, Redis et MinIO sont bien présents dans l'environnement Docker, mais la persistance métier a volontairement commencé par des fichiers JSON et PDF stockés localement pour accélérer la mise en service.

            Ce compromis n'affaiblit pas l'architecture. Au contraire, il montre une hiérarchisation pragmatique des chantiers : d'abord prouver la valeur du flux utilisateur, puis étendre la maturité des composants d'infrastructure autour de lui. Le DAT assume cette trajectoire progressive et la documente comme telle.
            """
        ).strip(),
    )

    add_heading(document, "2.3 Contraintes de démonstration et de sécurité", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Une contrainte structurante du projet est l'environnement de démonstration lui-même. La toolbox devait rester utilisable en localhost, mais aussi pouvoir exploiter des outils plus naturellement disponibles sur une machine Kali Linux. Cette contrainte a conduit à dissocier l'interface d'orchestration de l'environnement d'exécution. Le site web reste dans Docker, tandis que les commandes peuvent être expédiées en SSH vers une VM Kali. Cette séparation répond à la fois à une contrainte d'outillage et à un souci de propreté opérationnelle.

            Une seconde contrainte provenait du comportement de certains outils sur des cibles de laboratoire. Des essais trop agressifs pouvaient faire chuter Juice Shop ou saturer l'application conteneurisée. Cela a entraîné plusieurs décisions concrètes : limitation des threads de Gobuster, limitation de Hydra au seul protocole SSH, réduction des paramètres de SQLmap, temps de capture bornés pour Wireshark et distinction claire entre audit contrôlé et bruteforce.

            Enfin, la solution devait rester présentable. Ce point, souvent sous-estimé dans des projets purement techniques, a influencé la mise en page de l'interface, le nommage des scripts, l'organisation du code et la façon dont les rapports sont générés. Le DAT présente ces éléments comme des décisions d'architecture à part entière, car ils participent directement à la qualité perçue de la solution.
            """
        ).strip(),
    )

    add_requirement_matrix(document)

    add_heading(document, "4. Principes directeurs de l'architecture", 1)
    add_heading(document, "4.1 Principe de simplicité d'usage", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le premier principe directeur a été la simplicité d'usage. Une grande partie de la valeur du projet repose sur sa capacité à rendre accessibles des outils habituellement réservés à des utilisateurs familiers de la ligne de commande. Pour atteindre cet objectif, l'interface a été réduite à un modèle très lisible : une cible, une liste d'outils, une console, un historique. Les options avancées n'apparaissent que lorsqu'elles sont utiles, par exemple pour Hydra, Wireshark ou Auth Audit.

            Ce principe de simplicité ne signifie pas simplisme. Les traitements complexes n'ont pas été supprimés ; ils ont été déplacés dans le backend. Le travail de normalisation des cibles, de validation des options, de masquage de secrets, de construction de commandes et de gestion d'erreurs se fait côté service. L'interface n'est donc qu'une façade maîtrisée sur un ensemble de logiques métiers plus riches.
            """
        ).strip(),
    )

    add_heading(document, "4.2 Principe de modularité", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le second principe directeur est la modularité. Chaque outil est décrit par un descripteur de module et, dans la plupart des cas, par un service dédié responsable de sa normalisation, de la construction de la commande et de sa persistance. Ce modèle permet d'ajouter un nouveau module sans remettre en cause l'architecture générale de l'application. Il suffit d'ajouter une entrée dans le registre, un service adapté et les routes API correspondantes.

            Cette modularité est également visible dans le classement par pôles. Le registre ne retourne pas une simple liste d'outils ; il construit des groupes cohérents, chacun associé à un besoin métier. Le frontend réutilise directement ces groupes pour composer l'interface. La cohérence entre données métier, backend et interface constitue l'une des forces structurelles du projet.
            """
        ).strip(),
    )

    add_heading(document, "4.3 Principe de séparation des environnements", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le troisième principe directeur est la séparation entre l'environnement d'orchestration et l'environnement d'exécution. L'application web ne cherche pas à embarquer tous les outils de sécurité dans son propre conteneur d'exécution fonctionnelle. Elle peut le faire localement pour certains usages, mais la trajectoire dominante du projet est de déléguer les commandes à Kali Linux. Cette stratégie permet de profiter d'un environnement spécialisé, de limiter les conflits de dépendances et d'éviter que l'interface ne devienne un conteneur monolithique difficile à maintenir.

            Dans le DAT, ce principe est essentiel, car il justifie la présence d'une brique SSH au cœur du backend. La machine Kali n'est pas un accessoire de démonstration ; elle est une extension naturelle de l'architecture. Le site web joue le rôle de cockpit, et Kali joue le rôle d'environnement opérateur. Cette distinction clarifie aussi les responsabilités de sécurité et facilite l'évolution de la solution.
            """
        ).strip(),
    )

    add_placeholder(document, "[Schéma - Vue logique de l'architecture globale : navigateur, conteneurs, Kali, lab cible]")

    add_heading(document, "5. Architecture d'infrastructure et conteneurisation", 1)
    add_heading(document, "5.1 Rôle de Docker Compose", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Docker Compose a été choisi comme mécanisme de composition de l'environnement local. Ce choix répond à un besoin de reproductibilité et de démonstration : l'utilisateur peut démarrer rapidement l'interface, les services associés et, si besoin, le lab Juice Shop. L'usage de Compose permet également d'aligner le projet sur les recommandations du cahier des charges en matière de conteneurisation, sans nécessiter la complexité d'une plateforme d'orchestration lourde pour une démonstration locale.

            Le fichier de composition centralise la définition des services, des variables d'environnement, des dépendances de démarrage et des profils optionnels. Il expose un service web, un worker Celery, PostgreSQL, Redis, MinIO et Juice Shop. Même si tous ces composants ne sont pas encore sollicités par le flux métier principal, leur présence documente une architecture cible cohérente et rend les prochaines évolutions beaucoup plus simples à enclencher.
            """
        ).strip(),
    )

    add_heading(document, "5.2 Inventaire des services conteneurisés", 2)
    add_table(document, ["Service", "Fonction", "Description"], DOCKER_SERVICES, [2.2, 4.0, 9.8])
    add_placeholder(document, "[Capture - docker compose ps montrant la stack active]")

    add_heading(document, "5.3 Profil lab et environnement de test", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Un profil Compose dédié au lab permet de lancer OWASP Juice Shop en même temps que la toolbox. Cette intégration répond à un besoin très concret : disposer d'une cible applicative locale, connue et adaptée à des démonstrations de scans web. Juice Shop joue ici le rôle d'environnement d'expérimentation, notamment pour WhatWeb, Gobuster, Nikto, SQLmap et Auth Audit.

            Le projet a également été validé avec une VM Kali exécutée dans VMware, reliée à l'environnement via SSH. Dans cette configuration, Juice Shop peut être visé depuis Kali via l'adresse IP de l'hôte Docker. Cette topologie a été documentée précisément parce qu'elle conditionne l'usage réel des outils : le navigateur peut rester sur le poste local, la toolbox tourne en conteneur, et les scanners s'exécutent sur Kali.
            """
        ).strip(),
    )

    add_heading(document, "5.4 Lecture réseau de l'environnement", 2)
    add_paragraphs(
        document,
        dedent(
            """
            D'un point de vue réseau, l'architecture doit être lue comme un ensemble de plans distincts. Le navigateur de l'utilisateur dialogue avec le conteneur web via le port 8000 exposé sur l'hôte. Le conteneur web dialogue ensuite avec ses services annexes par le réseau Docker interne : PostgreSQL pour la persistance cible, Redis pour l'orchestration cible et MinIO pour le stockage objet cible. Lorsqu'une campagne de scan doit être exécutée sur Kali, le conteneur initie en plus une connexion SSH vers la machine virtuelle située sur le réseau de virtualisation.

            Cette lecture est importante car elle explique plusieurs comportements observés pendant le projet. Par exemple, la saisie de 127.0.0.1 depuis l'interface ne vise pas automatiquement la machine hôte, mais le conteneur lui-même. De même, la portée réelle d'un scan dépend du point de vue réseau de l'outil exécuté. Documenter ce sujet dans le DAT permet de justifier certaines consignes d'utilisation, comme l'usage de host.docker.internal ou l'emploi des adresses IP de l'hôte pour les labs web.
            """
        ).strip(),
    )

    add_heading(document, "6. Architecture applicative", 1)
    add_heading(document, "6.1 Initialisation de l'application FastAPI", 2)
    add_paragraphs(
        document,
        dedent(
            """
            L'application est initialisée dans un point d'entrée central qui configure le logging, prépare les répertoires de travail puis instancie FastAPI. Cette phase d'initialisation crée les répertoires storage, reports, logs et auth si nécessaire. Cette logique de bootstrap garantit qu'un démarrage sur un environnement neuf ne nécessite pas d'étapes manuelles supplémentaires pour créer l'arborescence de persistance locale.

            FastAPI a été retenu pour plusieurs raisons. D'abord, il apporte nativement une documentation OpenAPI exposée via /docs, ce qui répond directement à l'exigence de Swagger. Ensuite, il gère efficacement les routes HTTP et les flux de réponse, ce qui a facilité l'implémentation des endpoints de streaming. Enfin, il permet de conserver un backend moderne, lisible et facilement extensible.
            """
        ).strip(),
    )

    add_heading(document, "6.2 Couche web et rendu Jinja2", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La couche web repose sur Jinja2 pour le rendu des pages HTML. Le routeur web gère la page principale, la page de connexion, les pages de configuration et de vérification de la double authentification, ainsi que la sonde de santé. Ce choix est cohérent avec la nature locale et démonstrative de l'application : Jinja2 permet de produire une interface légère, maîtrisée et directement connectée aux données du backend sans introduire la complexité d'un framework front-end séparé.

            La page principale reçoit du backend l'ensemble des groupes de modules, les paramètres par défaut utiles à certains outils et les indicateurs liés à l'authentification. L'interface est donc majoritairement pilotée par la configuration serveur. Cette approche simplifie le maintien en cohérence entre ce qui est réellement disponible dans le backend et ce qui est proposé à l'utilisateur dans l'interface.
            """
        ).strip(),
    )

    add_heading(document, "6.3 Couche API et contrat de données", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La couche API s'appuie sur des schémas Pydantic standardisés. Les structures ScanRequest, ScanResponse, ModuleExecutionResult, SessionCreateRequest et SessionResponse définissent un contrat de données homogène entre le navigateur et le backend. Cette homogénéité simplifie la maintenance, améliore la lisibilité de la documentation Swagger et garantit que tous les modules remontent des données structurées de manière comparable.

            Le routeur API expose à la fois des endpoints unitaires par module et des endpoints de streaming. Les modules qui s'appuient sur des CLI externes héritent très souvent de la même classe de service de streaming, ce qui mutualise une grande partie du comportement d'exécution. Cette factorisation constitue une décision d'architecture majeure : elle réduit les duplications de code et limite le coût d'ajout de nouveaux modules.
            """
        ).strip(),
    )

    add_heading(document, "6.4 Frontend JavaScript et expérience temps réel", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le front-end repose sur un script JavaScript unique chargé de plusieurs responsabilités : gestion du thème clair/sombre, sélection des modules, affichage conditionnel des panneaux de configuration, ouverture du flux NDJSON, alimentation de la console, création de la session finale et chargement de l'historique. Cette approche n'a rien d'un front-end complexe, mais elle est pleinement cohérente avec le besoin local et les objectifs de sobriété du projet.

            L'un des points les plus importants de cette couche est la gestion de l'ordre d'exécution. Lorsqu'un utilisateur coche plusieurs modules, l'interface ne les lance pas en parallèle ; elle les exécute l'un après l'autre selon un ordre explicitement défini. Cette décision a permis d'éviter des collisions d'affichage, de garder une lecture linéaire de la console et de produire un rapport unique par session. Dans le contexte du projet, ce choix améliore considérablement la lisibilité.
            """
        ).strip(),
    )
    add_placeholder(document, "[Capture - Console temps réel et panneau historique sur la même page]")

    add_heading(document, "6.5 Contrat d'affichage et cohérence de l'interface", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le front-end n'est pas conçu comme un catalogue statique d'options. Il suit un contrat d'affichage aligné sur les modules réellement disponibles. La liste affichée à l'utilisateur provient du registre des modules, ce qui évite les décalages entre l'interface et le backend. Lorsqu'un module n'est pas enregistré, il ne peut pas apparaître à l'écran ; à l'inverse, lorsqu'un module est ajouté au registre et branché à l'API, son exposition dans la page principale peut être maintenue sous contrôle.

            Ce principe se prolonge dans l'affichage conditionnel des paramètres avancés. Hydra, Wireshark et Auth Audit disposent de panneaux dédiés qui ne s'ouvrent que si le module correspondant est sélectionné. Cette logique évite d'encombrer la page principale par des formulaires inutiles, tout en conservant la possibilité d'exposer des réglages spécifiques lorsque cela est nécessaire. Du point de vue de l'architecture produit, cette décision est essentielle : elle maintient la simplicité générale de l'outil sans sacrifier les cas d'usage plus riches.
            """
        ).strip(),
    )

    add_heading(document, "7. Mécanisme d'exécution des outils et délégation Kali", 1)
    add_heading(document, "7.1 Service d'exécution générique", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le moteur d'exécution central distingue deux modes : local et SSH. Dans les deux cas, il retourne le même type de résultat, composé d'une sortie standard, d'une sortie d'erreur, d'un code de retour et d'une durée. Cette abstraction permet aux couches supérieures de raisonner en termes de module et de session, sans avoir à connaître le mode de transport réellement utilisé.

            En mode local, le service s'appuie sur subprocess.Popen et lit les sorties via des threads dédiés afin de pouvoir streamer la commande au fil de l'eau. En mode SSH, il ouvre une connexion Paramiko vers Kali, exécute la commande en pseudo-terminal et récupère progressivement stdout et stderr depuis le canal. Dans les deux cas, une logique de timeout protège l'application contre les commandes bloquantes.
            """
        ).strip(),
    )

    add_heading(document, "7.2 Délégation vers Kali Linux", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La délégation des commandes vers Kali Linux constitue l'un des choix les plus structurants du projet. Plutôt que d'essayer de maintenir dans le conteneur web l'ensemble des outils de pentest, de leurs wordlists et de leurs dépendances, il a été jugé plus rationnel de s'appuyer sur une VM Kali déjà adaptée à cet usage. Cette décision réduit le coût de maintenance de la stack web et améliore la compatibilité avec des outils comme Hydra, Gobuster, Nikto ou theHarvester.

            Le mode SSH est piloté par des variables d'environnement dédiées : hôte, port, identifiant, mot de passe ou chemin de clé. Le service accepte l'auto-ajout des empreintes hôtes lorsque la configuration le permet, ce qui simplifie un environnement de lab. En cas d'erreur d'authentification ou d'indisponibilité réseau, un message explicite est renvoyé à l'interface. Cette lisibilité est importante, car elle rend les échecs compréhensibles pour l'utilisateur final.

            Il faut noter que la solution actuelle repose sur un compte administré manuellement côté Kali. Cette méthode est adaptée à une maquette ou à un laboratoire, mais elle devra être renforcée dans une trajectoire de production : clés dédiées, rotation des secrets, cloisonnement plus fin des droits et meilleure gouvernance des accès distants.
            """
        ).strip(),
    )

    add_heading(document, "7.3 Streaming NDJSON et session unifiée", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le streaming a été implémenté via des réponses NDJSON, chaque ligne portant un événement de type start, stdout, stderr, finish ou error. Ce protocole léger présente deux avantages. D'une part, il est facile à consommer en JavaScript sans infrastructure temps réel plus lourde. D'autre part, il s'adapte aussi bien à des modules CLI qu'à des modules applicatifs natifs comme Auth Audit.

            Une fois les modules exécutés, le frontend ne considère pas chaque résultat comme une fin en soi. Il construit une session globale, composée de la cible, de la liste ordonnée des modules et des résultats unitaires. Le backend agrège alors les sorties, les durées et les codes de retour pour créer un artefact de session unique. Ce mécanisme répond directement au besoin utilisateur formulé pendant le projet : n'avoir qu'une seule trace d'historique et qu'un seul jeu de rapports pour une campagne multi-outils.
            """
        ).strip(),
    )
    add_placeholder(document, "[Schéma - Séquence d'une session : navigateur, API, service, Kali, artefacts]")

    add_heading(document, "7.4 Normalisation des cibles et validation métier", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Un point souvent peu visible mais central dans l'architecture est la normalisation des cibles. L'utilisateur peut saisir une IP, un domaine ou une URL, mais tous les outils n'attendent pas le même format. Nmap et Hydra peuvent fonctionner sur une IP ou un nom d'hôte. SQLmap exige une URL complète. theHarvester a besoin d'un domaine et refuse une IP brute. Auth Audit exige une URL de page de connexion. La couche service de chaque module a donc pour responsabilité de traduire l'entrée utilisateur en une cible compatible avec son propre moteur.

            Ce travail de validation métier évite plusieurs catégories d'erreurs : erreurs de syntaxe en ligne de commande, exécutions inutiles, faux diagnostics et incompréhensions côté utilisateur. Il contribue également à la stabilité de l'interface, car l'erreur est captée tôt et restituée dans un message intelligible. Dans une application qui orchestre des outils hétérogènes, cette normalisation des entrées est un véritable composant d'architecture, pas un détail d'implémentation.
            """
        ).strip(),
    )

    add_heading(document, "8. Sécurité de la plateforme", 1)
    add_heading(document, "8.1 Authentification locale et A2F", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La plateforme propose une authentification locale optionnelle. Lorsqu'elle est activée, l'utilisateur ne peut plus accéder à la page principale tant que la session n'est pas authentifiée. Le premier facteur repose sur un identifiant et un mot de passe stockés dans l'environnement de l'application. Le second facteur repose sur un code TOTP à six chiffres, généré par une application de type Google Authenticator, Microsoft Authenticator ou Aegis.

            Le parcours d'authentification a été conçu en deux temps. Après validation du premier facteur, l'utilisateur est redirigé soit vers une page de configuration initiale de la double authentification, soit vers une page de vérification du code OTP si l'A2F est déjà activée. La configuration initiale affiche un QR code et la clé secrète TOTP. Une fois le code validé, l'état d'activation est persistant dans le répertoire storage/auth.

            Cette implémentation répond au besoin de sécuriser l'accès à la toolbox dans un contexte local partagé ou de démonstration. Elle montre également que l'équipe ne s'est pas contentée de sécuriser les scans eux-mêmes ; elle a pris en compte la sécurité du point d'entrée applicatif.

            Dans une version plus mature, cette logique gagnerait à être replacée dans un vrai modèle d'identité persistant. Une table users pourrait stocker l'identifiant, le hash de mot de passe, l'état du compte, le rôle et les dates de dernière activité. Une table auth_factors ou mfa_enrollments pourrait ensuite porter le secret TOTP chiffré, la date d'enrôlement, le dernier usage valide, l'état d'activation, ainsi qu'un mécanisme de rotation ou de réinitialisation. On pourrait y ajouter des codes de secours à usage unique, une invalidation administrative et, à terme, plusieurs facteurs par utilisateur. Cette évolution est importante, car elle conditionne le passage vers un vrai multi-utilisateur traçable et gouvernable.
            """
        ).strip(),
    )
    add_placeholder(document, "[Capture - Page de connexion puis page de vérification A2F]")

    add_heading(document, "8.2 Middleware de session et contrôle des routes", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le contrôle d'accès s'appuie sur SessionMiddleware. Une session authentifiée est matérialisée par des indicateurs simples placés dans request.session, dont authenticated et auth_stage. Les routes API sont protégées par une dépendance qui vérifie cet état avant d'autoriser l'accès. Ce mécanisme n'est pas sophistiqué, mais il est cohérent avec la portée actuelle du projet : un compte administrateur unique, une application locale et un besoin de protection simple mais réel.

            Le paramétrage actuel présente cependant des limites importantes qu'il faut documenter avec honnêteté. Le mot de passe et le secret TOTP sont aujourd'hui fournis en clair via les variables d'environnement. Le middleware de session est configuré avec same_site=lax et https_only=False, ce qui est acceptable pour du localhost mais ne constitue pas un niveau de durcissement suffisant pour une exposition plus large. Le futur DAT d'exploitation devra donc traiter ces points comme des recommandations prioritaires.

            Si l'application devait évoluer vers un mode réellement partagé, la session ne devrait plus être la seule source de vérité. Il serait préférable de rattacher chaque session à un user_id stocké en base, à un niveau d'authentification atteint et à une date d'expiration pilotable. Cette approche permettrait d'imposer une revalidation MFA pour certaines opérations sensibles, de fermer toutes les sessions d'un compte compromis, de distinguer proprement les droits lecture, exécution et administration, et d'aligner l'interface web avec d'éventuels appels API authentifiés.
            """
        ).strip(),
    )

    add_heading(document, "8.3 Journalisation d'audit", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La journalisation d'audit est assurée par un logger dédié. Chaque service de scan journalise son exécution en sérialisant un événement JSON dans storage/logs/audit.log. Les informations stockées incluent typiquement le nom du module, l'identifiant du scan, la cible, la durée et le code de retour. Ce mécanisme fournit une première capacité de traçabilité, utile aussi bien pour le support que pour la preuve d'usage.

            Le choix d'un RotatingFileHandler répond à un besoin de robustesse basique : les fichiers de logs n'augmentent pas sans limite et peuvent être archivés par rotation. Là encore, le dispositif n'a pas encore la richesse d'une véritable plateforme SIEM ou d'un pipeline de logs centralisé. Mais il introduit une habitude essentielle de sécurité : toute exécution de scan doit laisser une trace exploitable.
            """
        ).strip(),
    )

    add_heading(document, "8.4 Écarts de sécurité à traiter", 2)
    add_bullet_list(
        document,
        [
            "Absence actuelle de RBAC : la plateforme fonctionne avec un compte administrateur unique.",
            "Absence d'HTTPS applicatif dans le flux local courant.",
            "Secrets stockés dans les variables d'environnement et non chiffrés par Fernet dans le chemin métier actif.",
            "Connexion SSH vers Kali adaptée au lab, mais encore peu industrialisée pour un contexte multi-utilisateurs.",
            "Persistance des rapports sur disque local sans chiffrement ni politique d'archivage avancée.",
        ],
    )

    add_heading(document, "9. Stockage, reporting et persistance", 1)
    add_heading(document, "9.1 Génération des artefacts", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Les résultats des scans et des sessions sont persistés sous deux formes : JSON et PDF. Le format JSON porte la donnée complète et structurée. Il est exploitable pour un traitement machine ultérieur, pour une reprise dans un autre service ou pour une lecture détaillée. Le format PDF joue un rôle différent : il sert avant tout de restitution synthétique et présentable dans un contexte de soutenance, de partage ou d'archivage rapide.

            Le service d'artefacts opère une distinction importante entre scan unitaire et session multi-modules. Dans le premier cas, il sérialise directement un ScanResponse. Dans le second, il crée un SessionResponse global, agrège les commandes, concatène les sorties et ajoute la liste détaillée des résultats unitaires dans le JSON final. Ce modèle répond exactement à l'attente formulée pendant le projet : regrouper dans un seul objet les exécutions successives réalisées depuis l'interface.

            Le PDF actuel est produit par un moteur interne léger, fondé sur la construction directe d'un document PDF minimal. Ce choix permet de rester autonome et de ne pas dépendre d'un moteur externe plus lourd pour la restitution de base. En contrepartie, la mise en forme reste sobre. Dans une phase ultérieure d'industrialisation, il sera possible de remplacer ce moteur par un générateur plus riche, sans modifier le contrat fonctionnel attendu par l'interface.

            Une trajectoire réaliste consiste à mieux séparer la fabrication du contenu et son rendu final. Le JSON peut rester la source de vérité technique, puis alimenter un gabarit HTML ou DOCX plus riche avant conversion vers un PDF de qualité supérieure. Dans ce schéma, l'API OpenAI n'aurait pas vocation à "dessiner" le PDF elle-même, mais à enrichir certaines sections à forte valeur documentaire : synthèse exécutive, reformulation des résultats bruts, priorisation des risques, proposition de remédiations et homogénéisation du vocabulaire. Le rendu final, lui, resterait déterministe via un moteur de template et un convertisseur de document, ce qui préserverait la reproductibilité des livrables.

            Cette intégration doit toutefois être strictement encadrée. Les données transmises au modèle doivent être limitées au strict nécessaire, idéalement pseudonymisées lorsque le contexte l'exige, et chaque génération doit rester validable par un opérateur avant publication. L'IA peut accélérer la lecture, résumer et structurer, mais elle ne doit pas devenir une source opaque de décision ni inventer des vulnérabilités absentes du résultat brut. Le DAT a donc intérêt à présenter cette piste comme un enrichissement contrôlé du reporting, et non comme une substitution au traitement technique existant.
            """
        ).strip(),
    )
    add_placeholder(document, "[Capture - Rapport PDF d'une session comprenant plusieurs modules]")

    add_heading(document, "9.2 Historique et cycle de vie des rapports", 2)
    add_paragraphs(
        document,
        dedent(
            """
            L'historique visible dans l'interface interroge directement le répertoire des rapports JSON et reconstruit une vue synthétique des dernières exécutions. Chaque entrée contient l'identifiant de scan, le nom du module ou de la session, la cible, la commande, la durée, le code de retour et les liens de téléchargement vers les artefacts. La logique est volontairement simple et transparente.

            La contrepartie de cette simplicité est que la persistance métier reste aujourd'hui fondée sur le système de fichiers. Cette approche est pleinement acceptable pour une version de démonstration ou de laboratoire. Elle deviendra toutefois insuffisante dès lors que le projet visera un volume d'usage plus important, des recherches plus élaborées, une gestion de droits différenciés ou une gouvernance plus poussée des données. Le fait d'avoir déjà prévu PostgreSQL et MinIO dans l'environnement Docker limite néanmoins l'effort de transition futur.

            Le passage à PostgreSQL permettrait d'aller beaucoup plus loin sur l'historique. Une simple interface de recherche pourrait alors filtrer les sessions par utilisateur, cible, famille d'outils, plage de dates, statut d'exécution ou niveau de sévérité remonté dans le rapport. Ce changement améliorerait à la fois l'exploitation quotidienne, l'audit a posteriori et la capacité à construire des tableaux de bord de pilotage sans reparser manuellement des fichiers locaux.
            """
        ).strip(),
    )

    add_heading(document, "9.3 Position de PostgreSQL et MinIO dans l'architecture cible", 2)
    add_paragraphs(
        document,
        dedent(
            """
            PostgreSQL et MinIO ne sont pas des ajouts décoratifs. Ils matérialisent la direction cible du projet. PostgreSQL a vocation à devenir le référentiel des métadonnées de scans, des utilisateurs, des rôles, des historiques filtrables et des états applicatifs. MinIO a vocation à prendre en charge l'archivage des artefacts de session, des exports et potentiellement des pièces jointes futures liées à des campagnes.

            Le fait que ces deux briques ne soient pas encore au centre du flux métier courant doit être analysé comme un choix de séquencement, et non comme un manque de vision. L'équipe a d'abord sécurisé la chaîne de valeur visible par l'utilisateur. Le DAT documente donc une architecture où la couche de persistance avancée est prête, mais pas encore consommée par le cœur métier dans sa version actuelle.

            Concrètement, PostgreSQL pourrait être organisé autour de tables users, roles, scans, scan_results, sessions, artifacts, audit_events et saved_profiles. La base ne stockerait pas nécessairement les fichiers lourds eux-mêmes, mais plutôt leur identité, leur propriétaire, leur état, leur empreinte, leur localisation objet et leur politique de rétention. MinIO prendrait en charge les PDF, JSON, exports complémentaires et éventuelles pièces jointes, via un couple bucket/object_key versionnable. Cette séparation entre métadonnées relationnelles et artefacts objets est classique, lisible et parfaitement adaptée à une plateforme de ce type.
            """
        ).strip(),
    )

    add_heading(document, "9.4 Convention de nommage et cycle de vie des fichiers", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La persistance locale s'appuie sur une convention de nommage simple : chaque scan ou session reçoit un identifiant unique de type UUID, qui sert de base aux noms des artefacts JSON et PDF. Cette convention évite les collisions, facilite les téléchargements et permet de reconstruire l'historique sans base relationnelle active. Les répertoires storage/reports, storage/logs et storage/auth jouent ainsi un rôle de première couche documentaire.

            Cette stratégie, bien que simple, présente un intérêt fort dans une architecture de démonstration. Elle permet à un examinateur ou à un repreneur du projet d'ouvrir rapidement le contenu produit et de relier visuellement les rapports à une exécution. Le revers de cette lisibilité est l'absence actuelle de politique d'archivage, de purge ou de chiffrement local. Ces sujets devront être traités lors du passage à une persistance plus mature.

            Une évolution naturelle consisterait à introduire un cycle de vie explicite des artefacts : brouillon, validé, archivé, supprimé. Ce statut pourrait être porté par PostgreSQL tandis que MinIO gérerait la conservation physique, la rétention et le versioning. On gagnerait ainsi en gouvernance documentaire, en capacité de purge sélective et en conformité avec un usage plus durable que la simple démonstration locale.
            """
        ).strip(),
    )

    add_tool_sections(document)

    add_heading(document, "11. Maintenabilité, extensibilité et industrialisation", 1)
    add_heading(document, "11.1 Réutilisation par classes de service", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La plupart des modules de la toolbox reposent sur une classe commune de streaming de commandes. Ce point structurel est important pour la maintenabilité. Un nouveau module de type CLI n'a pas besoin de re-développer toute la chaîne de lecture des sorties, de sérialisation des événements, de génération des rapports et de persistance. Il doit avant tout fournir sa logique de normalisation, sa commande et ses messages métiers.

            Cette factorisation réduit les divergences d'implémentation entre modules. Elle explique aussi pourquoi l'application a pu intégrer progressivement de nouveaux outils sans se dégrader en accumulation de scripts spéciaux difficilement relisibles. Le DAT insiste sur cet aspect, car il s'agit d'un signe de maturité logicielle au-delà du simple ajout fonctionnel.
            """
        ).strip(),
    )

    add_heading(document, "11.2 Celery et Redis : état courant et potentiel", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Redis et Celery sont présents dans l'architecture comme briques d'industrialisation. Le worker Celery est défini dans Docker Compose et l'application dispose d'une configuration Celery minimale, avec sérialisation JSON et timezone UTC. Toutefois, le chemin d'exécution principal des scans reste synchrone. Ce choix a été assumé afin de privilégier la visibilité immédiate des sorties et la simplicité de démonstration.

            L'intérêt de Celery n'est pas remis en cause. Il deviendra particulièrement pertinent pour des scans longs, des campagnes planifiées, des traitements de masse, ou la génération différée de rapports enrichis. Redis pourra alors jouer pleinement son rôle de broker et de backend de résultats. Le DAT présente donc Celery comme une brique préparée et cohérente, mais encore non intégrée au scénario d'usage principal.
            """
        ).strip(),
    )

    add_heading(document, "11.3 Connecteur GitHub et logique d'extension", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le projet inclut un connecteur de base capable de récupérer un fichier brut depuis un dépôt GitHub. Cette brique est volontairement simple : elle construit une URL raw à partir d'un owner, d'un repository, d'une branche et d'un chemin, puis utilise httpx pour télécharger le contenu. Bien que cette intégration ne soit pas encore exposée dans l'interface principale, elle préfigure un besoin identifié très tôt : pouvoir brancher ultérieurement des scripts, listes ou templates maintenus dans un référentiel distant.

            Ce point illustre une logique d'architecture progressive. Tout n'est pas encore visible côté utilisateur, mais le backend est déjà préparé à absorber des cas d'usage supplémentaires sans rupture de structure. La capacité à consommer des ressources GitHub pourra notamment servir à alimenter des wordlists, des scripts personnalisés ou des modèles de paramétrage réutilisables.
            """
        ).strip(),
    )

    add_heading(document, "11.4 Gouvernance du code et lisibilité du projet", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La maintenabilité ne dépend pas uniquement d'un bon découpage logiciel ; elle dépend aussi de la propreté générale du dépôt. Dans ce projet, une attention particulière a été portée au nommage des fichiers et des services, à la séparation des responsabilités et à la limitation des scripts auxiliaires. Cette discipline répond à une demande explicite du projet : conserver un dossier présentable et compréhensible pour un tiers.

            Concrètement, les noms de modules, de services et de scripts ont été alignés sur leur fonction effective. Les routes sont regroupées, les services encapsulent la logique métier, les modules décrivent les capacités exposées, et les artefacts de démonstration sont stockés dans un espace dédié. Ce soin apporte une vraie valeur documentaire : le DAT peut s'appuyer directement sur l'arborescence sans devoir réinterpréter le projet à partir d'un ensemble hétérogène de fichiers.
            """
        ).strip(),
    )

    add_heading(document, "12. Exploitation et déploiement", 1)
    add_heading(document, "12.1 Démarrage de la plateforme", 2)
    add_paragraphs(
        document,
        dedent(
            """
            L'exploitation courante de la solution repose sur un cycle de démarrage simple. L'utilisateur prépare son fichier .env, construit l'image si nécessaire, puis lance docker compose up avec le profil adapté. En configuration standard, le service web suffit pour l'interface. En configuration enrichie, le profil lab ajoute Juice Shop et le profil worker active le worker Celery.

            Cette simplicité de démarrage constitue un atout important du projet. Elle permet de présenter la plateforme rapidement sans imposer une suite d'opérations manuelles complexes. Elle favorise également la reprise du projet par une autre personne, ce qui est un critère implicite mais réel de qualité d'architecture.
            """
        ).strip(),
    )
    add_bullet_list(
        document,
        [
            "Copie du modèle d'environnement : Copy-Item .env.example .env",
            "Démarrage standard : docker compose up --build web",
            "Démarrage avec le lab : docker compose --profile lab up --build",
            "Démarrage avec le worker : docker compose --profile worker up --build",
            "Accès principal : http://localhost:8000",
        ],
    )
    add_placeholder(document, "[Capture - Terminal montrant le lancement Docker Compose]")

    add_heading(document, "12.2 Scénario d'usage nominal", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Dans le scénario nominal, l'utilisateur se connecte à la toolbox, saisit une cible puis coche un ou plusieurs modules. Si un outil nécessite des paramètres complémentaires, les panneaux adéquats apparaissent dynamiquement sous la bonne catégorie. Une fois la campagne lancée, la console affiche la préparation, les commandes, les sorties et la fin de session. L'historique se met ensuite à jour avec une seule entrée agrégée.

            Ce scénario a été recherché dès les premières itérations du projet. Il répond à l'idée d'une plateforme de pilotage plutôt qu'à celle d'un simple lanceur de scripts. La navigation ne demande pas de connaissance avancée de la structure interne de l'application ; elle suit un parcours naturel centré sur la cible et les outils.
            """
        ).strip(),
    )

    add_heading(document, "12.3 Bonnes pratiques d'exploitation", 2)
    add_bullet_list(
        document,
        [
            "Privilégier l'usage du mode SSH Kali pour les outils dont l'écosystème est plus stable dans cette distribution.",
            "Réserver les scans les plus agressifs à des cibles de laboratoire ou à des environnements explicitement autorisés.",
            "Conserver les rapports de session comme trace d'exécution, en particulier pour les démonstrations et revues techniques.",
            "Mettre à jour régulièrement les wordlists et binaires côté Kali plutôt que de complexifier l'image web.",
            "Vérifier les paramètres de timeout lorsque de nouveaux modules sont ajoutés.",
        ],
    )

    add_heading(document, "12.4 Procédures de maintenance courante", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La maintenance courante de la plateforme se concentre aujourd'hui sur quatre axes. Le premier est la cohérence de l'environnement Docker : vérification de l'état des conteneurs, reconstruction des images lorsque le code évolue et contrôle des profils utilisés. Le second est la cohérence de la VM Kali : disponibilité SSH, présence des binaires attendus et actualisation des wordlists ou outils de test.

            Le troisième axe concerne les artefacts et les journaux. Le stockage local doit être surveillé, les rapports anciens peuvent être archivés si besoin, et le journal d'audit doit être consulté en cas d'anomalie ou de besoin de traçabilité. Le quatrième axe est purement applicatif : vérification que les écrans d'authentification, la configuration TOTP et le chargement des modules restent cohérents après toute évolution du code. Cette vision de maintenance a été intégrée au DAT pour montrer que la solution n'a pas seulement été codée ; elle a aussi été pensée pour être exploitée.
            """
        ).strip(),
    )

    add_heading(document, "12.5 Usage recommandé du lab Juice Shop", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Juice Shop occupe une place particulière dans le projet. Il ne constitue pas seulement un service additionnel ; il représente une cible de lab permettant de valider les modules applicatifs sans dépendre d'un site tiers. Son usage recommandé est clair : démonstrations de WhatWeb, Gobuster, Nikto, SQLmap et Auth Audit dans un cadre connu et isolé.

            Le DAT recommande toutefois de garder une approche mesurée lors des démonstrations. Même dans un lab prévu à cet effet, certains outils peuvent provoquer des lenteurs, voire des redémarrages du service si leurs paramètres sont trop agressifs. C'est précisément cette réalité qui a motivé le bridage de plusieurs modules. Documenter cette relation entre outil et cible de test renforce la crédibilité opérationnelle du projet.
            """
        ).strip(),
    )

    add_heading(document, "13. Scénarios de test et validation", 1)
    add_heading(document, "13.1 Stratégie de validation", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La validation de la plateforme n'a pas été limitée à un simple contrôle de démarrage. Chaque brique importante a été testée dans un scénario cohérent avec son usage. L'objectif n'était pas d'obtenir une couverture unitaire exhaustive, mais de vérifier le parcours complet : accès à l'interface, sélection des modules, streaming des sorties, agrégation des sessions, production des rapports et consultation de l'historique.

            La présence d'un lab Juice Shop et d'une VM Kali a fortement aidé cette validation, car elle a permis de tester la cohérence bout en bout du système dans un environnement proche de l'usage visé. La matrice suivante résume les principaux scénarios vérifiés.
            """
        ).strip(),
    )
    add_table(document, ["Scénario", "Objectif", "Statut", "Observation"], TEST_SCENARIOS, [4.0, 5.0, 1.8, 5.2])
    add_placeholder(document, "[Capture - Exemple de session multi-outils validée sur Juice Shop]")

    add_heading(document, "13.2 Retour d'expérience sur les essais", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Plusieurs essais ont conduit à des ajustements d'architecture ou d'intégration. Gobuster a dû être ralenti pour ne pas faire tomber Juice Shop ou geler l'interface. Hydra a été recentré sur SSH, car les tests de formulaires web modernes étaient trop instables et peu exploitables dans le cadre du projet. SQLmap a été simplifié afin de rester démontrable et de réduire les réponses interactives. Ces arbitrages montrent que l'équipe a piloté la solution à partir de l'observation réelle, et non à partir d'une simple liste de fonctionnalités imaginées.

            De la même manière, le passage par Kali Linux a représenté une évolution majeure du projet. Plutôt que de chercher à tout résoudre dans le conteneur web, l'architecture a été recentrée sur la bonne séparation des responsabilités. Ce retour d'expérience constitue un élément important du DAT, car il explique pourquoi l'architecture finale est plus crédible que la version initialement envisagée uniquement en local.
            """
        ).strip(),
    )

    add_heading(document, "13.3 Enseignements techniques principaux", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le premier enseignement technique du projet est qu'une bonne intégration d'outil passe autant par le cadrage de l'usage que par l'appel du binaire lui-même. Les modules les plus stables ne sont pas forcément ceux qui exposent le plus d'options, mais ceux qui limitent intelligemment l'espace de configuration pour produire un résultat prédictible.

            Le deuxième enseignement est que la séparation entre cockpit web et environnement opérateur spécialisé constitue une architecture particulièrement efficace pour ce type de projet. Elle permet de garder un front-end propre, maîtrisé et démontrable, tout en s'appuyant sur Kali Linux pour les outils les plus sensibles aux dépendances ou aux wordlists.

            Le troisième enseignement est que la traçabilité fait partie intégrante de la valeur d'une toolbox. Le live terminal, l'historique, l'agrégation des sessions et les rapports exportables sont au moins aussi importants que le scan lui-même. Sans eux, l'outil resterait un simple lanceur de commandes. Avec eux, il devient une plateforme de démonstration, de capitalisation et d'analyse.
            """
        ).strip(),
    )

    add_heading(document, "14. Risques, limites et feuille de route", 1)
    add_heading(document, "14.1 Limites actuelles", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La première limite de la solution réside dans son niveau d'industrialisation. L'application est parfaitement adaptée à un contexte de laboratoire, de démonstration ou de soutenance, mais plusieurs sujets restent ouverts pour un usage plus large : gestion multi-utilisateur, RBAC, chiffrement des secrets, stockage objet réel des artefacts, persistance relationnelle active et sécurisation HTTPS complète.

            Une seconde limite concerne le périmètre des outils. Certains scanners réputés du monde pentest n'ont pas encore été intégrés, comme OpenVAS, Metasploit ou ZAP. Leur absence ne remet pas en cause la validité du projet, mais elle limite encore la portée de la boîte à outils au regard du cadrage le plus ambitieux du cahier des charges. La modularité actuelle a toutefois été pensée pour faciliter leur intégration future.

            La troisième limite concerne la qualité de reporting. Les rapports PDF actuels sont fonctionnels, traçables et suffisants pour démontrer l'exécution, mais ils restent minimalistes sur le plan visuel. Un futur chantier pourrait viser une mise en page plus riche, une hiérarchisation plus poussée et des exports complémentaires comme CSV ou HTML dédiés.
            """
        ).strip(),
    )

    add_heading(document, "14.2 Principaux risques résiduels", 2)
    add_bullet_list(
        document,
        [
            "Mauvaise gouvernance des secrets si les variables d'environnement sont partagées sans contrôle.",
            "Résultats non exploitables si des outils trop agressifs sont lancés sur des cibles fragiles sans paramétrage adapté.",
            "Dépendance à la disponibilité de Kali en mode SSH pour une partie importante du périmètre fonctionnel.",
            "Dérive de la maintenabilité si de futurs modules sont ajoutés hors du patron de service commun.",
            "Écart entre architecture cible et architecture active si PostgreSQL, MinIO et Celery ne sont pas réellement branchés à moyen terme.",
        ],
    )

    add_heading(document, "14.3 Feuille de route recommandée", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La feuille de route recommandée peut être structurée en trois horizons. À court terme, il convient de consolider les briques déjà présentes : branchement de PostgreSQL pour l'historique, utilisation de MinIO pour les artefacts, amélioration du rendu PDF et durcissement de l'authentification locale. À moyen terme, l'objectif devrait être l'activation d'un mode asynchrone piloté par Celery, l'introduction d'un RBAC simple et l'ajout de nouveaux modules fortement attendus. À plus long terme, la plateforme pourrait évoluer vers un véritable socle d'orchestration multi-campagnes, avec plugin system, calendrier de scans et gouvernance plus fine.

            Cette feuille de route est réaliste parce qu'elle s'appuie sur des briques déjà préparées dans le projet. L'architecture n'a donc pas besoin d'être reconstruite pour évoluer ; elle a surtout besoin d'être approfondie et branchée plus complètement sur les services déjà prévus.
            """
        ).strip(),
    )

    add_heading(document, "14.4 Recommandations de durcissement prioritaire", 2)
    add_bullet_list(
        document,
        [
            "Remplacer le stockage des secrets en clair par une gestion chiffrée et une politique de rotation.",
            "Activer un schéma de rôles minimal avant tout élargissement du nombre d'utilisateurs.",
            "Basculer la persistance de l'historique vers PostgreSQL afin de faciliter l'audit, la recherche et la gouvernance.",
            "Externaliser les artefacts dans MinIO pour mieux isoler le stockage des rapports et préparer l'archivage.",
            "Introduire Celery pour les scans longs afin de dissocier plus proprement réactivité de l'interface et durée des traitements.",
            "Prévoir une terminaison TLS et une politique de journalisation enrichie en cas d'ouverture à un périmètre plus large que localhost.",
        ],
    )

    add_heading(document, "15. Conclusion", 1)
    add_paragraphs(
        document,
        dedent(
            """
            La Pentest Toolbox répond de manière concrète à un besoin de centralisation et de simplification de l'exécution d'outils de sécurité. Elle fournit un point d'entrée unique, une exécution pilotée, une restitution temps réel et un mécanisme d'historisation cohérent. Le projet ne se limite donc pas à juxtaposer des commandes ; il propose une architecture logicielle claire, extensible et déjà alignée sur plusieurs exigences du cahier des charges.

            Le principal intérêt de la solution réside dans la cohérence entre son objectif, son interface et son architecture. Le choix de FastAPI, de Jinja2, de Docker Compose, du mode SSH vers Kali et d'une logique de services par module aboutit à un ensemble lisible, compréhensible et maintenable. Cette cohérence est, en pratique, plus importante qu'une accumulation prématurée de fonctionnalités non stabilisées.

            Enfin, le projet présente une trajectoire crédible. Les briques nécessaires à sa montée en maturité sont déjà en place dans l'environnement : PostgreSQL, Redis, MinIO, Celery, documentation Swagger, modularité backend et connecteur GitHub. Le DAT ne présente donc pas seulement un état des lieux ; il formalise une base technique solide sur laquelle une version plus industrialisée pourra être construite sans rupture majeure de conception.
            """
        ).strip(),
    )

    add_heading(document, "16. Annexes techniques", 1)
    add_heading(document, "16.1 Endpoints API principaux", 2)
    add_table(document, ["Méthode", "Endpoint", "Usage"], API_ENDPOINTS, [2.0, 6.0, 8.0])

    add_heading(document, "16.2 Groupes de variables d'environnement", 2)
    rows = []
    for group, values in ENV_GROUPS:
        rows.append((group, "\n".join(values)))
    add_table(document, ["Groupe", "Variables concernées"], rows, [4.0, 12.0])

    add_heading(document, "16.3 Services et rôles associés", 2)
    add_table(document, ["Service", "Rôle principal", "Perspective d'évolution"], [
        ("web", "Interface et API", "Durcissement HTTPS, RBAC et optimisation de l'observabilité."),
        ("worker", "Exécution asynchrone préparée", "Branchement effectif des traitements longs."),
        ("postgres", "Persistance relationnelle cible", "Historique structuré, comptes utilisateurs, métadonnées."),
        ("redis", "Transport de tâches et résultats", "Montée en charge et planification de scans."),
        ("minio", "Stockage objet cible", "Archivage externalisé des artefacts et rapports."),
        ("juice-shop", "Lab de test", "Ajout d'autres labs spécialisés par usage."),
    ], [2.2, 5.0, 8.8])

    add_heading(document, "16.4 Liste indicative des captures à intégrer", 2)
    add_bullet_list(
        document,
        [
            "[Capture - Page de connexion]",
            "[Capture - Configuration A2F avec QR code]",
            "[Capture - Page principale de la toolbox]",
            "[Capture - Vue des modules classés par pôle]",
            "[Capture - Panneau Hydra affiché conditionnellement]",
            "[Capture - Panneau Auth Audit affiché conditionnellement]",
            "[Capture - Console temps réel pendant une session]",
            "[Capture - Historique avec liens JSON/PDF]",
            "[Capture - docker compose ps]",
            "[Capture - Session multi-outils sur Juice Shop]",
        ],
    )

    add_heading(document, "16.5 Glossaire fonctionnel et technique", 2)
    add_table(
        document,
        ["Terme", "Définition dans le contexte du projet"],
        [
            ("Session de scan", "Ensemble ordonné d'exécutions de modules lancé depuis une même action utilisateur et agrégé dans un rapport unique."),
            ("Module", "Capacité fonctionnelle exposée dans l'interface et appuyée par un service backend dédié."),
            ("Mode SSH", "Mode dans lequel la toolbox orchestre les outils depuis Docker mais exécute les commandes sur une VM Kali distante."),
            ("Artefact", "Fichier de sortie produit par la plateforme, typiquement en JSON ou PDF."),
            ("Auth Audit", "Module natif développé pour observer le comportement d'un parcours d'authentification sans réaliser de bruteforce massif."),
            ("Lab Juice Shop", "Application web vulnérable utilisée comme cible de démonstration et de validation des outils orientés web."),
        ],
        [3.5, 12.5],
    )

    save_document(document)


def tool_by_title(title: str) -> dict:
    try:
        return next(tool for tool in TOOLS if tool["title"] == title)
    except StopIteration:
        aliases = {
            "Hydra": "Hydra (mode SSH)",
        }
        alias = aliases.get(title)
        if alias:
            return next(tool for tool in TOOLS if tool["title"] == alias)
        raise


def add_tool_entry(document: Document, heading: str, tool_title: str, numbering: str):
    tool = tool_by_title(tool_title)
    full_heading = heading if heading.startswith(f"{numbering} ") else f"{numbering} {heading}"
    add_heading(document, full_heading, 3)
    p = document.add_paragraph()
    label = p.add_run("Commande type : ")
    ensure_font(label, bold=True)
    run = p.add_run(tool["command"])
    ensure_font(run)
    add_paragraphs(document, tool["body"])
    add_placeholder(document, tool["capture"])


def build_document_reconforge():
    document = Document()
    configure_section(document.sections[0])
    create_styles(document)
    add_header_footer(document.sections[0], first_page=True)

    add_cover(document)

    add_heading(document, "Table des matieres", 1)
    add_bullet_list(
        document,
        [
            "1. Introduction et contexte du projet",
            "2. Architecture generale de Reconforge",
            "3. Module Reconnaissance & Infrastructure",
            "4. Module OSINT & Fingerprinting",
            "5. Module Enumeration, Vulnérabilité & Contrôle Web",
            "6. Module Plateforme, Reporting & Orchestration",
            "7. Matrice de composants et interdépendances",
            "8. Planification projet sur 6 mois",
            "9. Conclusion et perspectives",
            "10. Annexes techniques",
            "Glossaire",
        ],
    )

    add_heading(document, "1. Introduction et contexte du projet", 1)
    add_heading(document, "1.1 Presentation generale", 2)
    add_paragraphs(
        document,
        dedent(
            f"""
            {APP_NAME} est une plateforme locale d'orchestration d'outils de pentest conçue pour simplifier, encadrer et tracer l'exécution de plusieurs utilitaires de sécurité depuis une interface web unique. Le projet répond à une problématique très concrète : les outils de cybersécurité offensive sont puissants, mais dispersés, hétérogènes, souvent difficiles à installer proprement et encore plus difficiles à démontrer dans un cadre pédagogique ou de soutenance.

            La réponse apportée par {APP_NAME} consiste à centraliser ces outils derrière une interface sobre et pilotable. L'utilisateur saisit une cible, choisit un ou plusieurs modules, suit la sortie des commandes en direct dans une console web, puis récupère un rapport unique de session en JSON et en PDF. L'intérêt principal de la solution n'est donc pas de remplacer les outils existants, mais de les rendre plus lisibles, plus cohérents et plus présentables dans un environnement de démonstration locale.

            Le projet a été construit à partir des technologies et des contraintes réellement mises en œuvre dans l'environnement partagé : Python 3.12, FastAPI, Jinja2, Docker Compose, PostgreSQL, Redis, MinIO, exécution locale ou distante via SSH, et intégration d'outils déjà présents ou facilement maintenables sur une VM Kali Linux. Le présent DAT décrit cette solution en s'appuyant sur ce qui existe effectivement, sans transformer le document en promesse théorique déconnectée de l'implémentation.
            """
        ).strip(),
    )

    add_heading(document, "1.2 Positionnement du projet et objectifs", 2)
    add_paragraphs(
        document,
        dedent(
            f"""
            Reconforge a été pensé comme un cockpit d'orchestration local. L'idée n'est pas de recréer un framework de pentest complet, ni d'empiler des outils derrière une même page web, mais de proposer un point d'entrée clair pour lancer, suivre et conserver des analyses de sécurité dans un cadre maîtrisé. Cette logique explique les grands choix de conception : une interface compacte, un découpage modulaire, une journalisation d'audit, une agrégation par session et une séparation nette entre la couche web et l'environnement Kali utilisé pour l'exécution.

            Le projet poursuit quatre objectifs complémentaires. D'abord, mettre à disposition une base réellement exploitable, capable de lancer plusieurs outils utiles dès la première version. Ensuite, structurer l'application de manière suffisamment propre pour faciliter l'ajout de nouveaux modules. Il s'agit aussi de produire un rendu crédible pour une démonstration ou une soutenance, avec une interface lisible et des rapports cohérents. Enfin, l'architecture doit rester ouverte à une montée en maturité vers plus d'automatisation, une meilleure persistance et un niveau de sécurité renforcé.
            """
        ).strip(),
    )

    add_heading(document, "1.3 Les trois surfaces de Reconforge", 2)
    add_paragraphs(
        document,
        dedent(
            f"""
            {APP_NAME} repose sur trois surfaces d'usage qui se complètent naturellement. La première est l'interface web locale. C'est elle que l'utilisateur manipule au quotidien pour saisir une cible, sélectionner les modules, suivre la console temps réel et consulter l'historique des sessions.

            La deuxième surface est l'API FastAPI documentée via Swagger. Elle structure les échanges entre le front-end et le backend, expose le catalogue des modules et fournit les endpoints utiles pour lancer les scans et récupérer les artefacts. Cette couche est importante, car elle montre que la plateforme ne se limite pas à une page HTML ; elle peut aussi être pilotée de manière standardisée.

            La troisième surface correspond aux artefacts produits à la fin d'une exécution : JSON et PDF. Ces fichiers jouent le rôle de mémoire opérationnelle. Ils conservent la trace d'une session, facilitent la relecture d'un test et servent de support de restitution. Ensemble, ces trois surfaces donnent à Reconforge une structure cohérente, lisible et facilement présentable.
            """
        ).strip(),
    )
    add_placeholder(document, "[Capture - Page d'accueil de Reconforge avec modules et console]")

    add_heading(document, "2. Architecture generale de Reconforge", 1)
    add_heading(document, "2.1 Vue d'ensemble", 2)
    add_paragraphs(
        document,
        dedent(
            f"""
            L'architecture générale de {APP_NAME} suit une logique en couches assez simple à lire. La présentation repose sur FastAPI et Jinja2. La logique métier est portée par les routeurs API, les services d'exécution et le registre des modules. L'exécution elle-même se fait soit localement, soit à distance par SSH vers une machine Kali Linux. Enfin, la persistance s'appuie aujourd'hui sur les artefacts stockés sur le système de fichiers, tout en préparant l'arrivée de PostgreSQL, Redis et MinIO pour les étapes suivantes.

            Cette organisation se retrouve directement dans l'arborescence du projet. Les routes, les services, les modules, les intégrations et les templates sont séparés de manière lisible. Ce découpage facilite la compréhension du code et limite le risque de dérive vers un backend monolithique. Il rend aussi l'ajout de nouvelles fonctionnalités plus naturel, car chaque responsabilité a déjà sa place.
            """
        ).strip(),
    )

    add_heading(document, "2.2 Technologies communes", 2)
    add_table(
        document,
        ["Technologie / composant", "Categorie", "Justification du choix"],
        [
            ("Python 3.12", "Langage principal", "Lisible, rapide à prototyper et très adapté à l'intégration d'outils de sécurité et de traitements HTTP."),
            ("FastAPI", "Framework web et API", "Documentation Swagger native, bonne gestion des routes et des flux, structure moderne pour un backend modulaire."),
            ("Jinja2", "Rendu HTML", "Interface légère, cohérence directe avec le backend, absence de complexité front-end inutile."),
            ("Docker Compose", "Conteneurisation", "Démarrage reproductible de l'ensemble de la stack locale et préparation d'une démonstration maîtrisée."),
            ("PostgreSQL", "Persistance cible", "Base relationnelle prévue pour structurer les métadonnées et l'historique applicatif."),
            ("Redis", "Broker / backend", "Brique préparée pour la montée en charge et l'orchestration asynchrone via Celery."),
            ("MinIO", "Stockage objet cible", "Base cohérente pour externaliser les rapports et artefacts dans une version plus mature."),
            ("Paramiko", "SSH Python", "Connexion vers Kali Linux sans sortir du backend applicatif."),
            ("Pydantic", "Schémas de données", "Contrat clair entre frontend, API et services."),
            ("httpx", "Client HTTP", "Base propre pour Auth Audit et le connecteur GitHub."),
        ],
        [3.6, 3.0, 11.4],
    )

    add_heading(document, "2.3 Flux de donnees", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le parcours principal reste volontairement linéaire. L'utilisateur saisit une cible, choisit un ou plusieurs modules, puis le front-end transmet la demande à l'API. Le service concerné normalise les entrées, construit la commande, lance l'exécution et renvoie la sortie en streaming vers le navigateur. En fin de session, les résultats sont regroupés, puis enregistrés en JSON et en PDF.

            Ce flux résume bien la valeur ajoutée de la plateforme. Reconforge ne se contente pas de déclencher un binaire ; elle encadre l'exécution, met en forme les paramètres, restitue la sortie en direct, agrège les résultats et conserve une trace exploitable. Autrement dit, une commande système devient ici une fonctionnalité applicative complète.
            """
        ).strip(),
    )
    add_placeholder(document, "[Schéma - Flux de données : interface -> API -> service -> Kali/outil -> artefacts]")

    add_heading(document, "3. Module Reconnaissance & Infrastructure", 1)
    add_heading(document, "3.1 Presentation du module", 2)
    add_heading(document, "3.1.1 Role dans la chaine d'audit", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Cette première famille de modules regroupe les fonctions de découverte technique et de contrôle d'infrastructure. Elle intervient tôt dans une démarche d'évaluation de sécurité, lorsqu'il faut identifier les services visibles, observer un flux réseau ou vérifier la robustesse d'un accès SSH. Dans Reconforge, ce bloc rassemble Nmap, Wireshark/tshark et Hydra en mode SSH. Malgré leurs usages différents, ces outils ont un point commun : ils travaillent au plus près de l'exposition technique de la cible.

            Les rassembler dans une même partie améliore la lisibilité d'ensemble. La lecture suit ainsi une progression assez naturelle : d'abord la reconnaissance et l'infrastructure, ensuite la collecte d'information et le fingerprinting, puis les contrôles web plus ciblés.
            """
        ).strip(),
    )
    add_heading(document, "3.1.2 Architecture technique du module", 3)
    add_paragraphs(
        document,
        dedent(
            """
            D'un point de vue technique, ce bloc repose sur le même socle de services que le reste de l'application. Chaque service valide ses entrées, prépare sa commande, encadre les erreurs puis délègue l'exécution au moteur commun, en local ou via SSH. Cette base partagée permet d'obtenir un comportement homogène, même lorsque les outils intégrés sont très différents.

            Chaque module conserve malgré tout ses particularités. Hydra est le seul de ce bloc à afficher des paramètres conditionnels dans l'interface. Wireshark adapte son timeout à la durée de capture demandée. Nmap, de son côté, limite strictement les scripts NSE autorisés. Ce mélange de cadre commun et d'ajustements ciblés donne un ensemble à la fois stable et souple.
            """
        ).strip(),
    )

    add_heading(document, "3.2 Outils integres et taches realisees", 2)
    add_tool_entry(document, "3.2.1 Nmap — Reconnaissance reseau maitrisee", "Nmap", "3.2.1")
    add_tool_entry(document, "3.2.2 Wireshark / tshark — Observation de flux", "Wireshark / tshark", "3.2.2")
    add_tool_entry(document, "3.2.3 Hydra — Tests d'authentification SSH", "Hydra", "3.2.3")

    add_heading(document, "3.3 Technologies et justifications techniques", 2)
    add_table(
        document,
        ["Technologie", "Categorie", "Justification approfondie"],
        [
            ("Nmap", "Scan reseau principal", "Outil central pour la reconnaissance active, facile à cadrer et particulièrement démonstratif dans une interface de pilotage."),
            ("tshark", "Capture reseau", "Version CLI de Wireshark adaptée à un usage backend et à une restitution textuelle synthétique."),
            ("Hydra", "Authentification infrastructure", "Permet de représenter des tests de robustesse d'accès tout en restant cantonné à SSH pour conserver un comportement fiable."),
            ("Paramiko", "Transport SSH", "Pont applicatif vers Kali Linux, évitant d'alourdir excessivement le conteneur web."),
        ],
        [3.0, 3.2, 11.8],
    )

    add_heading(document, "3.4 Etat d'integration du module", 2)
    add_table(
        document,
        ["Capacite", "Statut", "Observation technique"],
        [
            ("Nmap", "Termine", "Disponible avec scripts autorisés et reporting complet."),
            ("Wireshark / tshark", "Termine", "Disponible avec durée manuelle, interface configurable et synthèse statistique."),
            ("Hydra SSH", "Termine", "Disponible avec wordlists Kali et masquage du mot de passe manuel."),
            ("Autres outils d'infrastructure", "A venir", "Le patron d'intégration est prêt pour d'autres scanners ou validateurs réseau."),
        ],
        [4.0, 2.5, 11.5],
    )

    add_heading(document, "4. Module OSINT & Fingerprinting", 1)
    add_heading(document, "4.1 Presentation du module", 2)
    add_heading(document, "4.1.1 Positionnement fonctionnel", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Ce bloc regroupe les fonctions de qualification de surface externe et de collecte d'indices publics. Dans Reconforge, il s'appuie sur theHarvester et WhatWeb. Le premier travaille principalement à partir d'un nom de domaine et de sources ouvertes ; le second observe directement une cible HTTP pour en déduire les technologies visibles. L'association est logique : l'un aide à cartographier l'exposition, l'autre à qualifier rapidement la pile web apparente.

            Dans un scénario de pentest, cette étape vient généralement avant les scans applicatifs plus lourds. Elle sert à mieux comprendre la cible, son exposition publique et les indices exploitables disponibles en amont. La séparer clairement des scans actifs rend la lecture du document plus fluide et reflète mieux le déroulement réel d'une phase de reconnaissance.
            """
        ).strip(),
    )
    add_heading(document, "4.1.2 Logique technique d'integration", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Techniquement, ce module repose sur le même socle commun que les autres, mais avec des règles de validation plus métier. WhatWeb accepte une URL complète ou une cible normalisée. theHarvester, à l'inverse, attend un domaine et refuse une IP. Cette différence est prise en charge côté service, sans casser l'homogénéité de l'interface ni du reporting.

            C'est précisément là que l'architecture montre son intérêt : des outils de nature différente peuvent être intégrés sans donner l'impression d'un assemblage hétérogène. L'utilisateur garde une expérience simple, pendant que le backend absorbe les particularités de chaque cas.
            """
        ).strip(),
    )

    add_heading(document, "4.2 Outils integres et taches realisees", 2)
    add_tool_entry(document, "4.2.1 theHarvester — OSINT domaine et surface externe", "theHarvester", "4.2.1")
    add_tool_entry(document, "4.2.2 WhatWeb — Fingerprint de technologies web", "WhatWeb", "4.2.2")

    add_heading(document, "4.3 Technologies et justifications techniques", 2)
    add_table(
        document,
        ["Technologie", "Categorie", "Justification approfondie"],
        [
            ("theHarvester", "OSINT domaine", "Permet de remonter rapidement des informations de surface à partir de sources publiques structurées."),
            ("WhatWeb", "Fingerprint web", "Outil rapide, lisible et bien adapté à une première qualification de la pile applicative."),
            ("httpx", "Support de collecte", "Brique déjà présente côté application pour les usages web et les futures extensions d'intégration."),
        ],
        [3.2, 3.2, 11.6],
    )

    add_heading(document, "4.4 Etat d'integration du module", 2)
    add_table(
        document,
        ["Capacite", "Statut", "Observation technique"],
        [
            ("theHarvester", "Termine", "Sources par défaut limitées et adaptées à une démonstration raisonnable."),
            ("WhatWeb", "Termine", "Intégration simple, robuste et immédiatement exploitable."),
            ("Autres sources OSINT", "A venir", "Possibilité d'ajouter des outils complémentaires sans remettre en cause le socle."),
        ],
        [4.0, 2.5, 11.5],
    )

    add_heading(document, "5. Module Enumeration, Vulnerabilite & Controle Web", 1)
    add_heading(document, "5.1 Presentation du module", 2)
    add_heading(document, "5.1.1 Positionnement dans la chaine de test", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Ce bloc rassemble les outils qui vont plus loin dans l'analyse d'une cible web ou d'un mécanisme d'authentification. Il comprend Gobuster, SQLmap, Nikto, SSLyze et Auth Audit. Leurs usages sont différents, mais ils ont tous le même objectif général : observer le comportement d'un service exposé pour faire ressortir un contenu, une faiblesse, une mauvaise configuration ou un manque de robustesse.

            Ce regroupement donne une place claire aux fonctions d'énumération, de vulnérabilité et de contrôle web. Il reflète aussi fidèlement le périmètre réel de la plateforme : seules les capacités effectivement intégrées sont présentées, sans extrapolation.
            """
        ).strip(),
    )
    add_heading(document, "5.1.2 Contraintes d'exploitation observees", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Les essais menés pendant le projet ont montré que ce module est le plus sensible aux contraintes de stabilité du lab. C'est ici que plusieurs garde-fous ont été introduits : réduction des threads, délais sur Gobuster, simplification du profil SQLmap, limitation de Hydra au seul SSH, et choix d'un audit contrôlé plutôt qu'un bruteforce web massif pour Auth Audit. Ces ajustements ne sont pas des limitations arbitraires ; ils constituent une réponse d'architecture à des phénomènes concrets observés pendant les tests.
            """
        ).strip(),
    )

    add_heading(document, "5.2 Outils integres et taches realisees", 2)
    add_tool_entry(document, "5.2.1 Gobuster — Enumeration de contenu", "Gobuster", "5.2.1")
    add_tool_entry(document, "5.2.2 SQLmap — Verification d'injection", "SQLmap", "5.2.2")
    add_tool_entry(document, "5.2.3 Nikto — Audit web de configuration", "Nikto", "5.2.3")
    add_tool_entry(document, "5.2.4 SSLyze — Verification TLS", "SSLyze", "5.2.4")
    add_tool_entry(document, "5.2.5 Auth Audit — Controle du parcours d'authentification", "Auth Audit", "5.2.5")

    add_heading(document, "5.3 Technologies et justifications techniques", 2)
    add_table(
        document,
        ["Technologie", "Categorie", "Justification approfondie"],
        [
            ("Gobuster", "Enumeration web", "Permet de cartographier rapidement des contenus exposés, avec des garde-fous adaptés au lab."),
            ("SQLmap", "Analyse d'injection", "Référence connue, ici cadrée dans une exécution non interactive et à faible impact."),
            ("Nikto", "Vulnerabilite web", "Scanner simple à expliquer, utile pour les failles de configuration et l'exposition de contenu."),
            ("SSLyze", "Robustesse TLS", "Complète le périmètre par une lecture de la posture cryptographique."),
            ("Auth Audit", "Module natif web", "Démontre la capacité de Reconforge à embarquer des traitements applicatifs sans dépendre d'un binaire externe."),
        ],
        [3.2, 3.0, 11.8],
    )

    add_heading(document, "5.4 Etat d'integration du module", 2)
    add_table(
        document,
        ["Capacite", "Statut", "Observation technique"],
        [
            ("Gobuster", "Termine", "Bridé volontairement pour protéger Juice Shop et garder une démonstration fluide."),
            ("SQLmap", "Termine", "Profil standardisé, sans panneau complexe dans l'interface."),
            ("Nikto", "Termine", "Timeout élargi pour absorber une exécution plus longue."),
            ("SSLyze", "Termine", "Intégration simple, adaptée aux besoins de contrôle TLS."),
            ("Auth Audit", "Termine", "Module natif configurable, pensé pour une lecture responsable des parcours d'authentification."),
        ],
        [4.0, 2.5, 11.5],
    )

    add_heading(document, "6. Module Plateforme, Reporting & Orchestration", 1)
    add_heading(document, "6.1 Presentation du module", 2)
    add_heading(document, "6.1.1 Role transversal", 3)
    add_paragraphs(
        document,
        dedent(
            f"""
            Ce bloc transversal correspond au socle applicatif de {APP_NAME}. Il ne s'agit pas d'un outil de pentest au sens strict, mais de l'ensemble des briques qui rendent leur orchestration possible : FastAPI, Jinja2, JavaScript de streaming, historique, génération d'artefacts, authentification locale avec A2F, journalisation, Docker Compose, Redis, MinIO, PostgreSQL, worker Celery et connecteur GitHub préparatoire.

            C'est cette couche qui donne sa valeur d'ensemble à la plateforme. Sans elle, Reconforge ne serait qu'une collection de wrappers CLI. Avec elle, l'application devient un environnement cohérent de pilotage, de restitution et de conservation des résultats.
            """
        ).strip(),
    )
    add_heading(document, "6.1.2 Positionnement dans l'architecture Reconforge", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Ce bloc se situe à l'intersection entre l'utilisateur, les services métiers et l'environnement d'exécution. Il reçoit les intentions, sécurise l'accès, orchestre les actions, diffuse les sorties en temps réel, agrège les sessions et conserve les résultats. Il constitue la colonne vertébrale de l'application, et sa stabilité conditionne directement la qualité perçue de l'ensemble du projet.
            """
        ).strip(),
    )

    add_heading(document, "6.2 Infrastructure de la plateforme", 2)
    add_heading(document, "6.2.1 Interface web locale", 3)
    add_paragraphs(
        document,
        dedent(
            """
            L'interface web locale repose sur un template Jinja2 unique et un fichier JavaScript principal. Elle a été volontairement maintenue sobre : barre de saisie cible, groupes de modules par pôle, console temps réel, historique compact, thème clair/sombre et paramètres contextuels. Cette retenue visuelle participe à la lisibilité du produit et à sa valeur démonstrative.
            """
        ).strip(),
    )
    add_heading(document, "6.2.2 API FastAPI et endpoints de streaming", 3)
    add_paragraphs(
        document,
        dedent(
            """
            L'API FastAPI expose le catalogue des modules, les endpoints unitaires et les endpoints de streaming. Elle constitue un contrat structuré entre le navigateur et le backend. Le recours à des schémas Pydantic standardisés garantit une cohérence de données qui simplifie à la fois le développement et la documentation Swagger.
            """
        ).strip(),
    )
    add_heading(document, "6.2.3 Execution locale ou via Kali Linux", 3)
    add_paragraphs(
        document,
        dedent(
            """
            L'exécution peut se faire localement, mais la trajectoire principale du projet repose sur la délégation SSH vers Kali Linux. Ce choix répond à une logique opérationnelle : garder un conteneur web propre et léger, tout en profitant d'un environnement de sécurité déjà outillé pour les binaires, les wordlists et les dépendances spécifiques.
            """
        ).strip(),
    )
    add_heading(document, "6.2.4 Historique, rapports et audit", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Le service d'artefacts produit un JSON et un PDF pour chaque scan ou session. L'historique reconstruit ensuite une vue synthétique à partir de ces fichiers. En parallèle, un audit log dédié consigne les exécutions. Cette combinaison répond à un besoin central du projet : ne pas perdre la trace de ce qui a été lancé, pourquoi et avec quel résultat.

            Dans la version actuelle, cette logique reste volontairement simple et locale. Elle est très adaptée à une démonstration ou à un laboratoire, mais elle montre aussi sa limite dès que l'on imagine plusieurs opérateurs, des volumes plus élevés ou un besoin de filtrage fin. Une persistance relationnelle active permettrait par exemple de rechercher les sessions par utilisateur, cible, plage de dates, famille d'outils ou statut d'exécution, sans reparser des fichiers un par un.
            """
        ).strip(),
    )
    add_heading(document, "6.2.5 Authentification locale et TOTP", 3)
    add_paragraphs(
        document,
        dedent(
            """
            L'accès à la plateforme peut être protégé par un identifiant, un mot de passe et un second facteur TOTP. Le parcours de première connexion, la page de vérification OTP et l'état d'activation persistent dans le stockage local. Cette capacité matérialise un premier niveau de sécurité d'accès à l'outil lui-même.

            Pour une version plus industrialisée, cette mécanique gagnerait à s'appuyer sur une vraie base de données d'identité. Une table users pourrait porter l'identifiant, le hash du mot de passe, le rôle, l'état du compte et les dates de dernière activité. Une table auth_factors ou mfa_enrollments pourrait ensuite stocker le secret TOTP chiffré, la date d'enrôlement, le dernier usage valide, un indicateur de révocation et, si besoin, plusieurs facteurs par utilisateur. On pourrait également y associer des codes de secours à usage unique, une réinitialisation administrateur et une traçabilité complète des changements de facteur.

            L'intérêt n'est pas seulement théorique. En base, la double authentification devient gouvernable : fermeture des sessions d'un compte compromis, obligation de repasser le MFA après une action sensible, désactivation ciblée d'un facteur perdu, et distinction claire entre compte administrateur, opérateur et lecture seule. Cette trajectoire ferait passer Reconforge d'une protection locale crédible à un véritable socle multi-utilisateur.
            """
        ).strip(),
    )
    add_heading(document, "6.2.6 Connecteur GitHub et briques d'industrialisation", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Un connecteur GitHub minimal est déjà présent pour récupérer des fichiers bruts distants. En parallèle, Redis, MinIO, PostgreSQL et Celery sont déjà inclus dans l'environnement Compose. Toutes ces briques n'alimentent pas encore le flux principal, mais elles documentent clairement la trajectoire d'évolution et la volonté d'industrialisation du projet.

            PostgreSQL et MinIO sont particulièrement structurants pour la suite. Le premier a vocation à devenir la source de vérité pour les utilisateurs, les rôles, les scans, les sessions, les artefacts et les événements d'audit. Le second a vocation à stocker les fichiers produits, par exemple les PDF, JSON, exports complémentaires ou pièces jointes futures, tandis que la base relationnelle ne conserverait que leurs métadonnées, leur empreinte et leur cycle de vie. Cette séparation entre données de gestion et objets de restitution est propre, classique et adaptée à une plateforme d'orchestration.
            """
        ).strip(),
    )
    add_heading(document, "6.2.7 Detail de la connexion SSH vers Kali Linux", 3)
    add_paragraphs(
        document,
        dedent(
            """
            La connexion distante repose sur Paramiko et sur un jeu de variables d'environnement dédiées : hôte, port, nom d'utilisateur, mot de passe ou chemin de cle, ainsi qu'un indicateur de confiance de la cle d'hôte. Concrètement, lorsque le mode d'exécution SSH est actif, le backend n'essaie pas de lancer les binaires de sécurité dans son propre conteneur. Il ouvre une session distante vers Kali, exécute la commande demandée dans cet environnement et récupère progressivement stdout et stderr depuis le canal SSH.

            Cette mécanique présente plusieurs avantages. Elle évite d'alourdir inutilement l'image web avec des outils de pentest nombreux et parfois instables. Elle permet aussi de profiter directement des binaires, wordlists et dépendances déjà maintenus dans Kali Linux. Enfin, elle rend le projet plus crédible en exploitation réelle, car l'outil d'orchestration et l'environnement d'attaque restent clairement séparés. Cette séparation améliore aussi la lisibilité du DAT : l'interface pilote, Kali exécute, et les artefacts reviennent ensuite dans le flux applicatif.
            """
        ).strip(),
    )
    add_heading(document, "6.2.8 Cycle exact d'une commande", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Le cycle complet d'une commande suit toujours la même logique. L'utilisateur saisit d'abord une cible et sélectionne un module. Le frontend construit ensuite la requête adaptée et l'envoie à l'API. Le service associé commence par normaliser les entrées : domaine, IP, URL, paramètres propres au module, limites fonctionnelles et garde-fous. Une fois ces contrôles passés, la commande finale est assemblée dans le backend.

            À partir de ce point, le moteur commun d'exécution prend le relais. Il lance la commande en local ou via SSH, émet un événement de démarrage, puis stream la sortie en temps réel au format NDJSON. Le frontend affiche ces événements dans la console web et garde la chronologie complète de la session. À la fin, le backend récupère le code de retour, mesure la durée, consolide stdout et stderr, puis transmet l'ensemble au service d'artefacts pour produire le JSON et le PDF. Cette chaîne unifiée est importante : elle garantit que chaque outil suit le même contrat de sortie, même si sa logique interne diffère.
            """
        ).strip(),
    )

    add_heading(document, "6.3 Generation des rapports", 2)
    add_paragraphs(
        document,
        dedent(
            """
            La génération des rapports repose aujourd'hui sur deux formats : JSON et PDF. Le JSON conserve la donnée riche et structurée. Le PDF fournit une restitution lisible et partageable. L'agrégation par session a été un point clé du projet, car elle permet de faire correspondre le rapport à une campagne complète et non à une simple commande isolée. Cette logique rapproche Reconforge d'une plateforme d'orchestration plus mature.

            Une évolution très crédible consiste à séparer plus nettement la donnée source, la mise en forme et l'enrichissement documentaire. Le JSON peut rester le format canonique, puis alimenter un gabarit HTML ou DOCX plus riche avant conversion en PDF final. Dans ce schéma, l'API OpenAI ne servirait pas à "fabriquer le PDF" elle-même, mais à générer ou améliorer certaines sections : synthèse exécutive, reformulation des résultats techniques, hiérarchisation des risques, proposition de remédiations et harmonisation du ton du rapport. Le rendu final resterait quant à lui déterministe grâce à un moteur de template, ce qui préserverait la cohérence et la reproductibilité.

            Cette piste impose toutefois un cadre strict. Les données partagées avec le modèle doivent être limitées au strict nécessaire, idéalement pseudonymisées lorsque le contexte l'exige, et la sortie doit toujours être relue par un opérateur. L'IA apporte donc ici une valeur d'assistance documentaire et de gain de temps, pas une substitution à l'analyse technique brute ni une autorisation d'inventer des constats absents des résultats.
            """
        ).strip(),
    )
    add_heading(document, "6.3.1 Exemples de commandes effectivement orchestrees", 3)
    add_table(
        document,
        ["Module", "Contexte", "Commande representative", "Commentaire d'integration"],
        COMMAND_REFERENCE_ROWS,
        [2.6, 2.9, 7.4, 4.1],
    )
    add_paragraphs(
        document,
        dedent(
            """
            Les commandes ci-dessus ne sont pas de simples exemples théoriques. Elles correspondent au type de commandes réellement générées ou très proches des commandes produites par la plateforme une fois la cible, les paramètres et le contexte d'exécution stabilisés. Dans le DAT, cette vue est utile car elle montre que l'outil n'est pas une abstraction vide : il orchestre bien des binaires concrets, avec des garde-fous et une logique métier explicite.
            """
        ).strip(),
    )
    add_heading(document, "6.3.2 Structure des artefacts produits", 3)
    add_paragraphs(
        document,
        dedent(
            """
            L'artefact JSON conserve la structure détaillée de l'exécution : identifiant de scan ou de session, cible, liste des modules invoqués, commandes, code de retour, durée, flux standard, flux d'erreur et métadonnées complémentaires. Le PDF, lui, se concentre sur une restitution plus lisible : rappel de la cible, modules exécutés, chronologie synthétique, extraits de sortie utiles, code global et liens de cohérence entre les éléments de la session.

            Ce découplage entre format machine et format de lecture constitue un choix d'architecture à part entière. Le JSON est fait pour être réexploité, recherché ou sérialisé dans d'autres flux. Le PDF est fait pour être présenté, transmis ou relu rapidement. Cette dualité renforce la valeur du projet dans un contexte réel, car elle couvre à la fois le besoin technique de conservation et le besoin documentaire de restitution.

            Dans une version plus mature, on peut également faire évoluer le cycle de vie des artefacts eux-mêmes : brouillon, validé, archivé, supprimé. PostgreSQL porterait alors l'état documentaire, les droits d'accès, la date d'expiration et la traçabilité des validations, tandis que MinIO assurerait la conservation physique, le versioning et la rétention. Ce type de cycle de vie donne beaucoup plus de valeur à un rapport lorsqu'il doit être conservé, partagé ou gouverné dans la durée.
            """
        ).strip(),
    )
    add_placeholder(document, "[Capture - Historique Reconforge avec téléchargement JSON et PDF]")

    add_heading(document, "6.4 Technologies et justifications techniques", 2)
    add_table(
        document,
        ["Technologie", "Categorie", "Justification approfondie"],
        [
            ("FastAPI", "Backend / API", "Socle moderne, lisible et documenté, adapté au streaming et aux schémas de données structurés."),
            ("Jinja2", "Rendu HTML", "Excellent compromis entre sobriété, contrôle et simplicité de déploiement."),
            ("JavaScript vanilla", "Frontend léger", "Suffisant pour le streaming, l'historique et le thème sans surcharger la stack."),
            ("ArtifactService", "Reporting", "Point central de persistance et d'agrégation des sessions."),
            ("AuthService", "Sécurité d'accès", "Implémentation locale cohérente d'un login renforcé par TOTP."),
            ("Celery / Redis / MinIO / PostgreSQL", "Industrialisation", "Briques déjà prêtes pour les prochaines étapes de maturité."),
        ],
        [3.2, 3.0, 11.8],
    )

    add_heading(document, "7. Matrice de composants et interdependances", 1)
    add_heading(document, "7.1 Matrice des composants", 2)
    add_table(
        document,
        ["Composant", "Depend de", "Apporte a"],
        [
            ("Interface web", "Routeur web, registre des modules, JavaScript principal", "Expérience utilisateur et pilotage des scans"),
            ("Routeur API", "Schémas Pydantic, services de scan, service d'artefacts", "Accès programmatique et streaming"),
            ("Services de scan", "CommandExecutionService, configuration, service d'audit", "Exécution des outils et production des résultats"),
            ("CommandExecutionService", "subprocess ou Paramiko", "Mode local ou SSH vers Kali"),
            ("ArtifactService", "Système de fichiers local", "Historique, JSON, PDF, rapports de session"),
            ("AuthService", "Variables d'environnement, stockage auth local", "Protection d'accès et A2F TOTP"),
            ("Docker Compose", "Images et variables d'environnement", "Assemblage cohérent de la plateforme"),
        ],
        [3.5, 6.0, 8.5],
    )

    add_heading(document, "7.2 Interdependances techniques", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Les interdépendances de Reconforge sont principalement de trois ordres. Le premier est applicatif : le frontend dépend du registre des modules et des endpoints de streaming pour refléter fidèlement les capacités réelles du backend. Le deuxième est infrastructurel : la bonne exécution de certains modules dépend de la disponibilité de Kali Linux et du réseau permettant la connexion SSH. Le troisième est évolutif : les briques Redis, PostgreSQL, MinIO et Celery sont déjà présentes mais n'exprimeront pleinement leur valeur qu'une fois davantage branchées au flux métier.

            Cette matrice d'interdépendance est utile pour comprendre la maintenabilité du projet. Elle montre qu'une évolution sur un composant bien cadré, par exemple un nouveau module de scan, n'impose pas de réécriture de l'ensemble. En revanche, elle rappelle aussi que certaines évolutions de fond, comme la migration de la persistance vers PostgreSQL, devront être conduites de manière structurée pour ne pas casser la lecture simple et locale qui fait aujourd'hui la force de la plateforme.
            """
        ).strip(),
    )

    add_heading(document, "8. Planification projet sur 6 mois", 1)
    add_heading(document, "8.1 Organisation de l'equipe projet", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Dans une trajectoire de realisation concrete, Reconforge est envisage comme un projet de six mois, demarrant par un kick-off en semaine 1 puis deroule sur vingt-quatre semaines. Le cadrage retenu correspond a une petite mission de type prestation ou projet d'etude avance : une equipe coeur de deux personnes, appuyee par un commanditaire pedagogique pour les arbitrages et une relecture technique en phase finale.

            La repartition des roles a ete pensee pour rester lisible vis-a-vis d'un evaluateur ou d'un commanditaire externe. Gulraiz Mikail porte prioritairement les sujets backend, execution et securite technique, tandis qu'Arthur Flament prend en charge surtout l'interface, le reporting, la documentation, la recette et la mise en forme des livrables. Les sujets d'architecture, de planning et de soutenance restent toutefois traites de maniere transverse afin d'eviter un cloisonnement artificiel.
            """
        ).strip(),
    )
    add_table(
        document,
        ["Acteur", "Role projet", "Charge estimee", "Responsabilites principales"],
        PROJECT_ORG_ROWS,
        [3.2, 4.4, 2.0, 8.4],
    )

    add_heading(document, "8.2 Gouvernance, methode et backlog", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le pilotage retenu s'appuie sur une methode hybride. Le projet suit une logique proche de l'agile pour la construction de l'application, avec un backlog priorise, des increments demonstrables et des revues regulieres, mais conserve une discipline documentaire plus proche d'un projet forfaitaire classique pour le DAT, les annexes et la preparation du rendu final. Cette combinaison est adaptee a un projet technique devant etre a la fois code, explique et soutenu.

            Le backlog a ete structure autour de grands epics couvrant le socle d'infrastructure, les modules de scan, le reporting, l'authentification, la persistance, la documentation et la demonstration. Les chiffres ci-dessous ne pretendent pas reconstituer un historique exact de tickets reels ; ils constituent une valorisation realiste de pilotage, suffisante pour presenter le projet avec le niveau de maturite attendu d'un prestataire externe.
            """
        ).strip(),
    )
    add_table(
        document,
        ["Instance", "Cadence", "Participants", "Livrable ou objectif"],
        PROJECT_GOVERNANCE_ROWS,
        [2.6, 2.0, 4.3, 7.1],
    )
    add_table(
        document,
        ["Indicateur", "Volume", "Lecture projet"],
        PROJECT_BACKLOG_ROWS,
        [3.8, 2.0, 10.2],
    )

    add_heading(document, "8.3 Chiffrage et budget estimatif", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Afin de presenter le projet comme une mission professionalisee, le tableau suivant propose une valorisation indicative en jours.homme et en cout HT. Les montants sont volontairement credibles pour une petite prestation d'integration et d'industrialisation menee sur six mois, avec une equipe restreinte et un environnement de lab majoritairement open source. Ils ne correspondent pas a une facturation reelle, mais a un ordre de grandeur defendable en soutenance.

            Le chiffrage repose sur un volume valorise d'environ quatre-vingt-seize jours.homme pour la production principale, auquel s'ajoutent des couts forfaitaires de lab et une marge de pilotage pour les aleas. Cette lecture aide a objectiver la charge, a justifier la repartition des phases et a montrer que le projet a ete pense avec une logique de cout, de capacite et de priorisation.
            """
        ).strip(),
    )
    add_table(
        document,
        ["Poste", "Charge", "Hypothese de taux", "Cout estime HT"],
        PROJECT_BUDGET_ROWS,
        [5.6, 2.0, 3.0, 3.4],
    )
    add_paragraphs(
        document,
        "Le budget indicatif ainsi valorise s'etablit a 48 950 EUR HT. Dans un cadre academique, ce montant sert avant tout de repere de professionnalisation et de demonstration de pilotage, non de facture effective."
    )

    add_heading(document, "8.4 Logique de pilotage retenue", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le pilotage suit une progression logique. Le debut du projet est consacre au cadrage, au socle Docker, a la definition des regles d'execution et a la structuration du coeur applicatif. La phase centrale porte l'integration des modules, la stabilisation de l'execution via Kali Linux, l'amelioration de l'interface et la qualite du reporting. La fin de periode se concentre sur la persistance plus structuree, le durcissement, la documentation, la recette et la livraison finale.

            Cette logique de deroulement s'articule autour de quatre jalons principaux : validation du cadrage au kick-off, demonstration du socle technique, passage en beta fonctionnelle avec historique et artefacts, puis gel de livraison pour la soutenance. Chaque jalon dispose d'un livrable tangible, ce qui permet de montrer une trajectoire de projet concrete plutot qu'une simple accumulation de taches techniques.
            """
        ).strip(),
    )
    add_heading(document, "8.5 Vue synthétique janvier à juin", 2)
    add_table(
        document,
        ["Lot de travail", "Responsable principal", "Jan.", "Fev.", "Mars", "Avr.", "Mai", "Juin"],
        PROJECT_TIMELINE_ROWS,
        [5.0, 4.4, 1.1, 1.1, 1.2, 1.1, 1.1, 1.1],
    )
    add_paragraphs(
        document,
        dedent(
            """
            Dans ce tableau, la marque X signale les mois de mobilisation principale pour chaque lot. Il s'agit d'une vue de pilotage synthetique, equivalente a une lecture mini-Gantt, suffisante pour un DAT ou une soutenance. Elle permet de visualiser l'enchainement des grandes phases, les chevauchements utiles entre backend et frontend, ainsi que la concentration des efforts de finalisation sur la fin du projet.
            """
        ).strip(),
    )
    add_placeholder(document, "[Capture - Vue Jira ou planning projet janvier a juin]")

    add_heading(document, "9. Conclusion et perspectives", 1)
    add_heading(document, "9.1 Bilan du projet", 2)
    add_paragraphs(
        document,
        dedent(
            f"""
            {APP_NAME} apporte une réponse concrète, lisible et techniquement cohérente au besoin exprimé au départ. La solution permet de lancer plusieurs outils depuis une seule interface, de suivre leur exécution, de consolider les résultats et de les restituer dans des rapports exploitables. Elle ne prétend pas être un framework de sécurité exhaustif, mais elle constitue une base sérieuse, démontrable et extensible.

            Le principal mérite de la plateforme réside dans l'équilibre qu'elle atteint entre simplicité et rigueur. Les outils sont encapsulés, mais pas dénaturés ; l'interface est accessible, mais pas trompeuse ; l'architecture est sobre, mais préparée à évoluer. Cette cohérence donne à la solution une crédibilité forte dans le cadre d'un DAT.
            """
        ).strip(),
    )
    add_heading(document, "9.2 Perspectives d'evolution", 2)
    add_bullet_list(
        document,
        [
            "Branchement effectif de PostgreSQL pour l'historique structuré et les comptes applicatifs.",
            "Externalisation des artefacts dans MinIO.",
            "Activation de Celery pour les scans longs et les campagnes planifiées.",
            "Mise en place d'un RBAC minimal et d'un chiffrement plus robuste des secrets.",
            "Ajout d'un pipeline de reporting enrichi avec synthèse assistée et validation humaine avant export.",
            "Ajout de nouveaux outils compatibles avec la logique modulaire existante.",
        ],
    )
    add_heading(document, "9.2.1 Evolutions ergonomiques envisagees", 3)
    add_table(
        document,
        ["Axe", "Orientation proposee"],
        ERGONOMIC_ENHANCEMENT_ROWS,
        [4.4, 11.6],
    )
    add_heading(document, "9.2.2 Evolutions de securite prioritaires", 3)
    add_table(
        document,
        ["Domaine", "Etat actuel", "Evolution recommandee"],
        SECURITY_ENHANCEMENT_ROWS,
        [3.4, 6.1, 6.5],
    )
    add_heading(document, "9.2.3 Evolutions techniques et d'industrialisation", 3)
    add_paragraphs(
        document,
        dedent(
            """
            Au-delà des briques déjà identifiées, plusieurs évolutions techniques peuvent faire franchir un cap au projet. La première consiste à mieux séparer les profils d'exécution : mode local simplifié pour démonstration rapide, mode Kali SSH pour lab complet, puis mode asynchrone pour campagnes plus longues. La deuxième concerne la persistance : PostgreSQL peut devenir la source de vérité pour les utilisateurs, les sessions, les scans et les états de campagne, tandis que MinIO prendrait en charge la conservation longue durée des artefacts et pièces jointes volumineuses.

            Une autre piste importante concerne l'identité et le contrôle d'accès. Le passage à une vraie base relationnelle permettrait de sortir d'une logique à compte unique pour aller vers des comptes nominatifs, des rôles, une A2F enrôlée par utilisateur, des codes de secours et une révocation traçable. Cette étape est essentielle si la plateforme doit un jour être utilisée par plusieurs opérateurs ou dans un cadre plus encadré qu'un simple laboratoire local.

            Le reporting peut également franchir un cap qualitatif important. Une chaîne composée d'un JSON canonique, d'un moteur de templates, d'un rendu PDF plus soigné et d'une synthèse assistée par API OpenAI permettrait de produire des rapports plus lisibles, plus homogènes et mieux adaptés à différents publics. La bonne approche consiste toutefois à conserver une validation humaine finale et à limiter l'IA à un rôle de reformulation, de hiérarchisation et d'aide à la remédiation, jamais à l'invention du fond technique.

            Enfin, l'introduction d'un vrai système de plugins reste structurante. L'architecture actuelle est déjà modulaire, mais un plugin system formalisé permettrait de décrire les capacités d'un outil, ses formulaires, ses garde-fous, ses dépendances et sa stratégie de reporting sans toucher au cœur applicatif. La plateforme pourrait alors évoluer vers un pilotage plus riche des campagnes : sauvegarde de profils de scans, planification d'exécutions, comparaison entre sessions et vision consolidée multi-cibles. Ces ajouts ne changeraient pas seulement l'apparence du produit ; ils modifieraient son niveau de maturité opérationnelle.
            """
        ).strip(),
    )
    add_heading(document, "9.3 Cadre legal, ethique et bonnes pratiques", 2)
    add_paragraphs(
        document,
        dedent(
            """
            L'usage de Reconforge doit rester strictement encadré. Les outils intégrés à la plateforme peuvent produire des effets réels sur des cibles, qu'il s'agisse de découverte de services, de scans applicatifs ou de tests d'authentification. À ce titre, leur exécution doit être réservée à des environnements de laboratoire, à des systèmes explicitement autorisés ou à des périmètres cadrés dans un contexte académique ou professionnel légitime.

            Cette dimension légale et éthique fait partie intégrante de l'architecture du produit. Le choix d'une authentification d'accès, la limitation de certains modules, l'agrégation par session et la journalisation d'audit participent tous à une utilisation plus responsable et plus traçable de la plateforme.
            """
        ).strip(),
    )

    add_heading(document, "10. Annexes techniques", 1)
    add_heading(document, "10.1 Endpoints API principaux", 2)
    add_table(document, ["Methode", "Endpoint", "Usage"], API_ENDPOINTS, [2.0, 6.0, 8.0])

    add_heading(document, "10.2 Variables d'environnement principales", 2)
    rows = [(group, "\n".join(values)) for group, values in ENV_GROUPS]
    add_table(document, ["Groupe", "Variables concernees"], rows, [4.0, 12.0])

    add_heading(document, "10.3 Scenarios de test validates", 2)
    add_table(document, ["Scenario", "Objectif", "Statut", "Observation"], TEST_SCENARIOS, [4.0, 5.0, 1.8, 5.2])

    add_heading(document, "10.4 Liste indicative des captures a integrer", 2)
    add_bullet_list(
        document,
        [
            "[Capture - Page de connexion Reconforge]",
            "[Capture - Configuration A2F avec QR code]",
            "[Capture - Accueil Reconforge avec saisie de cible]",
            "[Capture - Modules classes par pole]",
            "[Capture - Console temps reel pendant une session]",
            "[Capture - Historique avec JSON et PDF]",
            "[Capture - docker compose ps]",
            "[Capture - Session multi-outils sur Juice Shop]",
            "[Schéma - Architecture globale de Reconforge]",
            "[Schéma - Flux de données et session de scan]",
        ],
    )
    add_heading(document, "10.5 Detail du parcours d'execution d'un scan", 2)
    add_paragraphs(
        document,
        dedent(
            """
            Le parcours d'exécution peut être résumé comme une chaîne de traitement complète. Étape 1 : l'utilisateur choisit une cible et un ou plusieurs modules. Étape 2 : le frontend transmet les paramètres au backend, qui applique la normalisation adaptée au module. Étape 3 : le service compose la commande réelle et l'envoie soit au moteur local, soit à la couche SSH. Étape 4 : la sortie standard et la sortie d'erreur sont streamées vers la console web. Étape 5 : la fin de commande déclenche l'agrégation et la génération des artefacts. Étape 6 : l'historique rend immédiatement ces artefacts téléchargeables.

            Cette décomposition mérite d'être explicite dans le DAT, car elle montre que la valeur du projet ne repose pas uniquement sur les binaires intégrés. Elle repose aussi sur la qualité du chaînage entre saisie utilisateur, validation métier, exécution contrôlée, restitution temps réel, historisation et reporting.
            """
        ).strip(),
    )
    add_heading(document, "10.6 Points de vigilance d'exploitation", 2)
    add_bullet_list(
        document,
        [
            "Verifier en priorite la disponibilite reseau entre le conteneur web et la VM Kali avant toute demonstration.",
            "Conserver des parametres de scan modérés sur les cibles de lab pour eviter les ralentissements ou redemarrages.",
            "Verifier l'existence et les droits des wordlists Kali pour Hydra et Gobuster.",
            "Ne pas reutiliser les secrets d'authentification locale ou SSH en clair dans un environnement expose.",
            "Conserver la logique de validation stricte des cibles afin d'eviter les commandes incoherentes ou dangereuses.",
        ],
    )

    add_heading(document, "Glossaire", 1)
    add_table(
        document,
        ["Terme", "Definition dans le contexte du projet"],
        [
            ("Reconforge", "Nom de la plateforme d'orchestration de scans et d'outils de pentest."),
            ("Session de scan", "Ensemble ordonné d'exécutions de modules lancé depuis une même action utilisateur."),
            ("Module", "Capacité fonctionnelle exposée dans l'interface et branchée au backend."),
            ("Mode SSH", "Mode dans lequel Reconforge orchestre les outils depuis Docker mais les exécute sur Kali Linux."),
            ("Artefact", "Fichier produit par la plateforme, généralement en JSON ou PDF."),
            ("Auth Audit", "Module natif qui analyse le comportement d'un parcours d'authentification web de manière contrôlée."),
            ("Lab Juice Shop", "Application web vulnérable utilisée comme cible locale de validation des modules web."),
        ],
        [3.5, 12.5],
    )

    save_document(document)


if __name__ == "__main__":
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build_document_reconforge()
